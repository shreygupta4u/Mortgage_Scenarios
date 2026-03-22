"""
mortgage_wireframe.py — Word (.docx) wireframe export for Canadian Mortgage Analyzer
Generates a layout-reference document with NO financial numbers or PII.
"""
import io
from datetime import date


def generate_wireframe_docx(base_info: dict = None) -> bytes:
    """Return bytes of a .docx wireframe document."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        import subprocess, sys
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "python-docx"],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
        )
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

    doc = Document()

    # ── Helpers ────────────────────────────────────────────────────────────────
    def set_cell_bg(cell, color_hex: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color_hex)
        tcPr.append(shd)

    def box(label: str, color: str = "E8F0FE"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[ {label} ]")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x88)
        run.font.bold = True
        from docx.oxml import OxmlElement as OE
        pPr = p._p.get_or_add_pPr()
        pBdr = OE("w:pBdr")
        for side in ["top", "bottom", "left", "right"]:
            b = OE(f"w:{side}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "4")
            b.set(qn("w:space"), "1")
            b.set(qn("w:color"), "3355AA")
            pBdr.append(b)
        pPr.append(pBdr)
        p.paragraph_format.space_after = Pt(6)

    def section_title(t: str):
        p = doc.add_heading(t, level=2)
        p.runs[0].font.color.rgb = RGBColor(0x1a, 0x3c, 0x5e)

    def field_row(labels: list):
        tbl = doc.add_table(rows=1, cols=len(labels))
        tbl.style = "Table Grid"
        for i, lbl in enumerate(labels):
            cell = tbl.rows[0].cells[i]
            cell.text = lbl
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_bg(cell, "EEF2FF")
        doc.add_paragraph()

    def metric_row(labels: list):
        tbl = doc.add_table(rows=2, cols=len(labels))
        tbl.style = "Table Grid"
        for i, lbl in enumerate(labels):
            cell = tbl.rows[0].cells[i]
            cell.text = lbl
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            set_cell_bg(cell, "D5E8F0")
            val_cell = tbl.rows[1].cells[i]
            val_cell.text = "████"
            val_cell.paragraphs[0].runs[0].font.size = Pt(10)
            set_cell_bg(val_cell, "F5F8FF")
        doc.add_paragraph()

    # ── Title page ─────────────────────────────────────────────────────────────
    h = doc.add_heading("🏠 Canadian Mortgage Analyzer", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("App Wireframe — Structure reference (no sensitive data)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph(f"Generated: {date.today().strftime('%B %d, %Y')}")
    doc.add_page_break()

    # ── Tab sections ───────────────────────────────────────────────────────────
    tab_fns = [
        ("Tab 1 — Setup & Overview",        _wf_setup),
        ("Tab 2 — Rate Change Scenarios",    _wf_scenarios),
        ("Tab 3 — Amortization Schedule",    _wf_schedule),
        ("Tab 4 — Prepayment Analysis",      _wf_prepayment),
        ("Tab 5 — Break Penalty",            _wf_breakpen),
        ("Tab 6 — Scenario Comparison",      _wf_comparison),
    ]
    for tab_title, fn in tab_fns:
        doc.add_heading(tab_title, level=1)
        fn(doc, box, section_title, field_row, metric_row)
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Per-tab wireframe sections ─────────────────────────────────────────────────

def _wf_setup(doc, box, st_, fr, mr):
    st_("Section A: Purchase & Down Payment")
    fr(["Purchase Price ($)", "Down Payment (%)", "Down Amount ($)", "CMHC Premium"])
    box("CMHC advisory info bar (insured / not insured / warning)")
    st_("Section B: Initial Mortgage Terms")
    fr(["Mortgage Type", "Payment Frequency", "Interest Rate (%)", "Amortization (yrs)"])
    fr(["Term (yrs)", "Mortgage Start Date", "", ""])
    st_("Section C: Additional Renewal Terms (collapsable expander)")
    box("➕ Add Past Renewal  |  [Renewal #1 row ▼]  |  [Renewal #2 row ▼]")
    st_("Section D: Past Prepayments Already Made (collapsable expander)")
    box("➕ Add Past Prepayment  |  [Prepayment #1 row ▼]")
    st_("Key Metrics at a Glance (9 tiles in 3 columns)")
    mr(["Initial Principal", "Balance @ Term End", "Original Amort Period"])
    mr(["Balance Today", "Principal Paid to Date", "Interest Paid to Date"])
    mr(["Remaining Amortization (yrs + end date)", "Total Interest", "Current Monthly Payment"])
    st_("Charts")
    box("Donut Chart — Principal vs Total Interest", color="FFF0E0")
    box("Stacked Bar — Yearly P & I (Past / Current / Post segments)", color="E0F0FF")
    st_("Actions")
    box("💾 Save Setup to DB  |  📄 Export App Wireframe (.docx)")


def _wf_scenarios(doc, box, st_, fr, mr):
    box("➕ New Scenario button")
    st_("Scenario expander (one per scenario)")
    fr(["Scenario Name", "Description (text area)"])
    box("🚀 Quick Templates checkbox → template selector + Apply")
    box("➕ Add Renewal Entry button")
    st_("Renewal row")
    fr(["Mode: By Date / By Period", "Effective Date", "Type: Fixed/Variable",
        "Rate (%)", "🗑️ Delete"])
    fr(["Term (years) selectbox"])
    box("Term start → end date caption")
    box("⚡ Early Renewal warning banner (if applicable)")
    fr(["Original Posted Rate (%)", "Current Posted Rate (%)"])
    box("Penalty radio: 3-Month Interest | IRD | Custom value + text box")
    fr(["Misc Fees ($)"])
    box("Penalty / break-even summary bar")
    st_("Variable Sub-scenarios (if Variable type selected)")
    box("➕ Add Sub-Scenario  |  [Sub-scenario name + rate changes]")
    st_("Scenario Results — Parallel Base vs Scenario metrics")
    mr(["Base Interest ($)", "Adjusted Interest ($)",
        "Current Remaining (yrs)", "Adjusted Remaining (yrs)",
        "Required Monthly Payment ($)"])
    st_("Payment & Amortization Impact")
    fr(["Monthly Payment ($) — editable number input"])
    mr(["Adjusted Remaining (yrs)", "Mortgage-Free By (date)"])
    box("Stacked Bar — P & I (Past / Current / Post), year x-axis", color="E0F0FF")
    box("Rate over time chart (date x-axis)")
    box("💾 Save scenario  |  Name input")
    doc.add_heading("Saved Scenarios (DB) section", level=3)
    box("Expander per saved scenario: 4 metrics | ✏️ Load for Editing | 🗑️ Delete")


def _wf_schedule(doc, box, st_, fr, mr):
    fr(["Scenario dropdown (Current Setup / Saved Scenarios)"])
    fr(["View Mode selectbox", "Highlight current month checkbox"])
    st_("Hierarchy View Mode (default)")
    box("◀ PAST  [collapsed]  | Total N payments")
    box("★ CURRENT (last 4 months + today)  [expanded]  | Balance: $████")
    box("    └─ Payment rows with yellow highlight on today row")
    box("▶ FUTURE  [collapsed]  | N remaining payments (blue-italic rows)")
    st_("Balance Chart")
    box("Line chart: Balance + Cumulative Interest vs Date  |  'Today' vline", color="E0F0FF")
    box("⬇️ Download CSV button")


def _wf_prepayment(doc, box, st_, fr, mr):
    fr(["Rate scenario base dropdown"])
    st_("Left Column — Annual Lump-Sum Prepayments")
    fr(["Annual Lump-Sum ($)", "Month each year", "Starting year", "# years"])
    fr(["Lender limit (%) slider", "", "", ""])
    st_("Right Column — Increased Regular Payments & One-Time")
    fr(["Increase type radio: Fixed $ | % increase | None"])
    fr(["One-time mode: By Date | By Period", "Date / Period #", "Amount ($)"])
    st_("Metrics")
    mr(["Interest (rate sc.)", "Interest (+ prepayments)", "Remaining (rate sc.)",
        "Remaining (+ prepayments)", "Total New Prepaid", "Interest ROI"])
    box("Stacked Bar — P & I with prepayment impact", color="E0F0FF")
    box("💾 Save Prepayment Scenario  |  Name input")


def _wf_breakpen(doc, box, st_, fr, mr):
    st_("Left Column — Inputs")
    fr(["Outstanding Balance ($)", "Contract Rate (%)",
        "Mortgage Type", "Months Remaining in Term"])
    fr(["Misc Fees ($)", "", "", ""])
    st_("Right Column — IRD Inputs")
    fr(["Posted Rate at Origination (%)", "Current Posted Rate for Remaining Term (%)"])
    st_("Penalty Calculation")
    box("Radio: 3-Month Interest ($████) | IRD ($████) | Custom value [____]")
    mr(["3 Months Interest", "IRD", "Penalty Applied", "Total Exit Cost"])
    st_("Break-even Analysis")
    fr(["New rate if you break (%) — slider"])
    mr(["Interest (Stay)", "Interest + Fees (Break)", "Net Savings"])
    box("Line chart: Net Savings vs New Rate  |  break-even line  |  current rate marker",
        color="E0F0FF")
    box("Monthly payment change + months to recoup info bar")


def _wf_comparison(doc, box, st_, fr, mr):
    box("Number of scenarios to compare: 2 / 3 / 4 radio")
    st_("Scenario columns (repeated per scenario)")
    fr(["Label", "Rate (%)", "Amort (yrs)", "Frequency", "Annual Lump ($)"])
    box("Note: Scenario 1 defaults to Current Mortgage (base)")
    st_("Results Table")
    mr(["Scenario", "Rate", "Amort", "Freq",
        "Orig Payment", "Current Payment", "Remaining", "Total Interest", "Total Paid"])
    box("Balance comparison chart — date x-axis, one line per scenario", color="E0F0FF")
    box("Stacked Bar P & I by scenario and year", color="E0F0FF")
    box("🏆 Best scenario callout  |  📥 Download Prompt (.txt)")
