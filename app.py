"""
Canadian Mortgage Analyzer — Streamlit App v7
12 changes:
1.  Scenario stacked bar chart restored to combined yearly chart (matching Image 2)
2.  Setup page split: Purchase/Down | Mortgage Terms | Additional Renewals (collapsable)
3.  Key metrics restructured: 9 specific metrics with tooltips
4.  Removed "Original Payoff" — replaced with clearer remaining amortization + end date
5.  Amortization schedule ends at actual remaining amortization
6.  Tab order: Setup | Rate Change Scenarios | Amortization Schedule | Prepayment | Break Penalty | Comparison
7.  "Base Remaining" on scenario page = current remaining amortization
8.  Single source of truth for scenarios — DB only, no local duplicates
9.  Comparison tab includes current mortgage + saved scenarios
10. Word wireframe export button (no sensitive data)
11. Chart annotations use margin/offset to avoid legend overlap
12. Amortization schedule hierarchy toggle: Past / Current / Future > Year > Date
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from datetime import date
from dateutil.relativedelta import relativedelta
import json, math, uuid, io

# ── PAGE CONFIG ───────────────────────────────────────────────────
st.set_page_config(page_title="🏠 Canadian Mortgage Analyzer",
                   page_icon="🏠", layout="wide",
                   initial_sidebar_state="collapsed")

st.markdown("""
<style>
.main-header{font-size:2.1rem;font-weight:700;color:#1a3c5e;margin-bottom:.1rem;}
.sub-header {font-size:.9rem;color:#555;margin-bottom:1.2rem;}
.mc {background:#f0f4f8;border-radius:9px;padding:.65rem 1rem;
     border-left:4px solid #1a3c5e;margin-bottom:.35rem;}
.mc h3{margin:0;font-size:.72rem;color:#666;text-transform:uppercase;letter-spacing:.05em;}
.mc p {margin:0;font-size:1.2rem;font-weight:700;color:#1a3c5e;}
.mc-g{border-left:4px solid #27ae60;}.mc-g p{color:#27ae60;}
.mc-r{border-left:4px solid #e74c3c;}.mc-r p{color:#c0392b;}
.mc-b{border-left:4px solid #2980b9;}.mc-b p{color:#2980b9;}
.warn{background:#fff3cd;border:1px solid #ffc107;border-radius:7px;padding:.6rem .9rem;margin:3px 0;}
.ok  {background:#d4edda;border:1px solid #28a745;border-radius:7px;padding:.6rem .9rem;margin:3px 0;}
.inf {background:#cce5ff;border:1px solid #004085;border-radius:7px;padding:.6rem .9rem;margin:3px 0;}
.pen {background:#f8d7da;border:1px solid #f5c6cb;border-radius:7px;padding:.6rem .9rem;margin:3px 0;}
.db-gate{max-width:460px;margin:70px auto;padding:2rem;
         background:#f8fafc;border:1px solid #d0d7de;border-radius:12px;
         box-shadow:0 4px 20px rgba(0,0,0,.08);}
div[data-testid="stDataFrame"] thead tr th{
    position:sticky!important;top:0;z-index:10;
    background:#1a3c5e!important;color:#fff!important;font-size:.8rem;padding:6px 8px;}
.seg-head{font-weight:700;font-size:.95rem;padding:.4rem .8rem;
          border-radius:6px;margin:.5rem 0 .2rem;}
</style>""", unsafe_allow_html=True)

# ──────────────────────────────────────────────────────────────────
# HELPERS
# ──────────────────────────────────────────────────────────────────
def _vline_x(v):
    try:
        import pandas as _pd
        return int(_pd.Timestamp(v).timestamp()*1000)
    except Exception: return v

def _year_of(d):
    if hasattr(d,"year"): return d.year
    try:
        import pandas as _pd; return _pd.Timestamp(d).year
    except Exception: return 0

# ──────────────────────────────────────────────────────────────────
# DATABASE
# ──────────────────────────────────────────────────────────────────
def get_db_connection(server, database, trusted, user="", pwd=""):
    try:
        import pyodbc
        cs=(f"DRIVER={{ODBC Driver 17 for SQL Server}};SERVER={server};DATABASE={database};"
            +("Trusted_Connection=yes;" if trusted else f"UID={user};PWD={pwd};"))
        conn=pyodbc.connect(cs,timeout=5); _init_db(conn); return conn,None
    except Exception as e: return None,str(e)

def _init_db(conn):
    c=conn.cursor()
    c.execute("""IF NOT EXISTS(SELECT*FROM sysobjects WHERE name='mortgage_setup' AND xtype='U')
      CREATE TABLE mortgage_setup(id INT IDENTITY PRIMARY KEY,
        saved_at DATETIME DEFAULT GETDATE(),setup_data NVARCHAR(MAX))""")
    c.execute("""IF NOT EXISTS(SELECT*FROM sysobjects WHERE name='mortgage_scenarios' AND xtype='U')
      CREATE TABLE mortgage_scenarios(id INT IDENTITY PRIMARY KEY,name NVARCHAR(200),
        created_at DATETIME DEFAULT GETDATE(),params NVARCHAR(MAX),summary NVARCHAR(MAX))""")
    conn.commit()

def db_load_setup(conn):
    if not conn: return None
    try:
        c=conn.cursor()
        c.execute("SELECT TOP 1 setup_data FROM mortgage_setup ORDER BY id DESC")
        r=c.fetchone(); return json.loads(r[0]) if r else None
    except Exception: return None

def db_save_setup(conn, data):
    if not conn: return False
    try:
        c=conn.cursor()
        c.execute("DELETE FROM mortgage_setup")
        c.execute("INSERT INTO mortgage_setup(setup_data) VALUES(?)",json.dumps(data,default=str))
        conn.commit(); return True
    except Exception: return False

def db_save_scenario(conn, name, params, summary):
    if not conn: return False
    try:
        conn.cursor().execute(
            "INSERT INTO mortgage_scenarios(name,params,summary)VALUES(?,?,?)",
            name,json.dumps(params),json.dumps(summary))
        conn.commit(); return True
    except Exception: return False

def db_update_scenario(conn, sid, name, params, summary):
    if not conn: return False
    try:
        conn.cursor().execute(
            "UPDATE mortgage_scenarios SET name=?,params=?,summary=? WHERE id=?",
            name,json.dumps(params),json.dumps(summary),sid)
        conn.commit(); return True
    except Exception: return False

def db_load_scenarios(conn):
    if not conn: return []
    try:
        c=conn.cursor()
        c.execute("SELECT id,name,created_at,params,summary FROM mortgage_scenarios ORDER BY created_at DESC")
        return [{"id":r[0],"name":r[1],"created_at":str(r[2]),
                 "params":json.loads(r[3]),"summary":json.loads(r[4])} for r in c.fetchall()]
    except Exception: return []

def db_delete_scenario(conn, sid):
    if not conn: return
    try: conn.cursor().execute("DELETE FROM mortgage_scenarios WHERE id=?",sid); conn.commit()
    except Exception: pass

# ──────────────────────────────────────────────────────────────────
# MATH
# ──────────────────────────────────────────────────────────────────
FREQ={"Monthly":{"n":12,"accel":False},"Semi-Monthly":{"n":24,"accel":False},
      "Bi-Weekly":{"n":26,"accel":False},"Accelerated Bi-Weekly":{"n":26,"accel":True},
      "Weekly":{"n":52,"accel":False},"Accelerated Weekly":{"n":52,"accel":True}}

def periodic_rate(annual_pct,n): return ((1+annual_pct/200)**2)**(1/n)-1

def calc_pmt(principal,annual_pct,n,amort_years,accel=False):
    if annual_pct==0:
        t=amort_years*n; return principal/t if t else 0
    r=periodic_rate(annual_pct,n); np_=amort_years*n
    pmt=principal*r*(1+r)**np_/((1+r)**np_-1)
    if accel:
        rm=periodic_rate(annual_pct,12); nm=amort_years*12
        pmt=(principal*rm*(1+rm)**nm/((1+rm)**nm-1))/(n/12)
    return pmt

def cmhc_premium(price,down):
    dp=down/price*100
    if dp>=20: return 0.0,0.0
    if price>1_500_000 or dp<5: return None,None
    ins=price-down; rt=0.04 if dp<10 else (0.031 if dp<15 else 0.028)
    p=ins*rt; return p,p*0.13

def date_to_period(td,sd,n):
    if isinstance(td,str): td=date.fromisoformat(td)
    if isinstance(sd,str): sd=date.fromisoformat(sd)
    return max(1,int(round((td-sd).days/365.25*n)))

def period_to_date(period,sd,n):
    if isinstance(sd,str): sd=date.fromisoformat(sd)
    if n==12: return sd+relativedelta(months=int(period-1))
    if n==24: return sd+relativedelta(days=int((period-1)*15))
    if n==26: return sd+relativedelta(weeks=int((period-1)*2))
    return sd+relativedelta(weeks=int(period-1))

def calc_remaining_years(balance,rate_pct,n,payment):
    if payment<=0 or balance<=0: return 0.0
    r=periodic_rate(rate_pct,n)
    if r==0: return balance/(payment*n)
    denom=payment-balance*r
    if denom<=0.01: return 999.0
    return math.log(payment/denom)/math.log(1+r)/n

def build_amortization(principal,annual_pct,n,amort_years,
                       accel=False,start_date=None,
                       extra_payments=None,rate_changes=None,term_periods=None):
    if start_date is None: start_date=date.today().replace(day=1)
    if isinstance(start_date,str): start_date=date.fromisoformat(start_date)
    pmt=calc_pmt(principal,annual_pct,n,amort_years,accel)
    r=periodic_rate(annual_pct,n)
    em={}
    if extra_payments:
        for ep in extra_payments: em[int(ep["period"])]=em.get(int(ep["period"]),0)+float(ep["amount"])
    rm={}
    if rate_changes:
        for rc in rate_changes: rm[int(rc["period"])]=float(rc["new_rate"])
    rows=[]; bal=float(principal); tp=amort_years*n
    cr=float(annual_pct); cur_r=r; ci=cp=cprep=0.0; pd_=start_date
    for i in range(1,int(tp)+1):
        if bal<=0.005: break
        if term_periods and i>term_periods: break
        if i in rm:
            cr=rm[i]; cur_r=periodic_rate(cr,n)
            pmt=calc_pmt(bal,cr,n,(tp-i+1)/n,accel)
        int_c=bal*cur_r; princ=min(max(pmt-int_c,0),bal)
        extra=min(em.get(i,0),max(bal-princ,0))
        bal-=princ+extra; ci+=int_c; cp+=pmt; cprep+=extra
        rows.append({"Period":i,"Date":pd_,"Year":((i-1)//n)+1,"CalYear":_year_of(pd_),
                     "Payment":round(pmt,2),"Interest":round(int_c,2),
                     "Principal":round(princ,2),"Prepayment":round(extra,2),
                     "Total Paid":round(pmt+extra,2),"Balance":round(max(bal,0),2),
                     "Rate (%)":round(cr,3),"Cum Interest":round(ci,2),
                     "Cum Principal":round(cp-ci,2),"Cum Prepayment":round(cprep,2)})
        if n==12: pd_+=relativedelta(months=1)
        elif n==24: pd_+=relativedelta(days=15)
        elif n==26: pd_+=relativedelta(weeks=2)
        else: pd_+=relativedelta(weeks=1)
    df=pd.DataFrame(rows)
    if df.empty: return df,{}
    ti=df["Interest"].sum(); tprep=df["Prepayment"].sum(); tt=df["Payment"].sum()+tprep
    last_pmt=float(df["Payment"].iloc[-1])
    return df,{"payment":round(last_pmt,2),"total_paid":round(tt,2),
               "total_interest":round(ti,2),"total_principal":round(df["Principal"].sum(),2),
               "total_prepaid":round(tprep,2),"end_balance":round(df["Balance"].iloc[-1],2),
               "payoff_periods":len(df),"payoff_years":round(len(df)/n,2),
               "interest_pct":round(ti/tt*100,1) if tt else 0}

def get_today_metrics(df,n):
    today=date.today()
    if df.empty: return {}
    def tod(v): return v if isinstance(v,date) else (v.date() if hasattr(v,"date") else date.fromisoformat(str(v)[:10]))
    past=df[df["Date"].apply(tod)<=today]
    if past.empty:
        return {"balance_today":float(df.iloc[0]["Balance"]),"principal_paid_today":0.0,
                "interest_paid_today":0.0,"period_today":0,
                "remaining_periods":len(df),
                "remaining_years":round(len(df)/n,1),"as_of_date":today.strftime("%b %d, %Y")}
    row=past.iloc[-1]
    remaining_periods=len(df)-int(row["Period"])
    remaining_years=round(remaining_periods/n,1)
    remaining_end=today+relativedelta(years=int(remaining_years),
                                       months=int((remaining_years%1)*12))
    return {"balance_today":float(row["Balance"]),
            "principal_paid_today":float(row["Cum Principal"]),
            "interest_paid_today":float(row["Cum Interest"]),
            "period_today":int(row["Period"]),
            "remaining_periods":remaining_periods,
            "remaining_years":remaining_years,
            "remaining_end_date":remaining_end.strftime("%b %Y"),
            "as_of_date":tod(row["Date"]).strftime("%b %d, %Y")}

def calc_break_penalty(bal,rate,mtype,orig_p,curr_p,months_left):
    mr=periodic_rate(rate,12); tmo=bal*mr*3
    if mtype=="Variable":
        return {"3_months_interest":round(tmo,2),"ird":None,"calc_penalty":round(tmo,2),"method":"3 months interest (variable)"}
    ird=max(bal*(orig_p-curr_p)/100*months_left/12,0); pen=max(tmo,ird)
    return {"3_months_interest":round(tmo,2),"ird":round(ird,2),"calc_penalty":round(pen,2),
            "method":"IRD" if ird>tmo else "3 months interest"}

# ──────────────────────────────────────────────────────────────────
# CHART: Combined stacked bar — matching Image 2 style
# FIX #1 & #11: single chart, annotations pushed above legend
# ──────────────────────────────────────────────────────────────────
def stacked_bar_pi(df, today_p, term_end_p, title="Principal & Interest"):
    """Combined yearly stacked bar with 3 color-coded segments.
    Past=grey, Current term=blue/red, Post-term=faded. Annotations above chart."""
    df2=df.copy()
    df2["Seg"]=np.where(df2["Period"]<=today_p,"past",
                 np.where(df2["Period"]<=term_end_p,"current","post"))
    colours={"past":("#888888","#bbbbbb"),"current":("#1a3c5e","#e74c3c"),"post":("#6a8dab","#d89090")}
    opacities={"past":0.8,"current":1.0,"post":0.55}
    seg_names={"past":"Past","current":"Current Term","post":"Post-term (projected)"}
    fig=go.Figure()
    for seg in ["past","current","post"]:
        g=df2[df2["Seg"]==seg].groupby("CalYear").agg(P=("Principal","sum"),I=("Interest","sum")).reset_index()
        if g.empty: continue
        cp,ci_=colours[seg]; op=opacities[seg]; nm=seg_names[seg]
        fig.add_bar(x=g["CalYear"].astype(str),y=g["P"],name=f"Principal — {nm}",
                    marker_color=cp,opacity=op,legendgroup=f"p_{seg}",
                    text=g["P"].apply(lambda v: f"${v/1000:.0f}k"),textposition="inside",textfont_size=8)
        fig.add_bar(x=g["CalYear"].astype(str),y=g["I"],name=f"Interest — {nm}",
                    marker_color=ci_,opacity=op,legendgroup=f"i_{seg}",
                    text=g["I"].apply(lambda v: f"${v/1000:.0f}k"),textposition="inside",textfont_size=8)
    # FIX #11: annotations at top of chart (y=1.12+) to stay above legend row
    if today_p>0:
        yr_row=df2[df2["Period"]==today_p]["CalYear"]
        if not yr_row.empty:
            fig.add_annotation(x=str(yr_row.iloc[0]),y=1.12,xref="x",yref="paper",
                               text="▼ Today",showarrow=False,font=dict(color="#27ae60",size=10,family="Arial Bold"),
                               xanchor="center")
    if 0<term_end_p<len(df2):
        te_df=df2[df2["Period"]==min(term_end_p,len(df2)-1)]
        if not te_df.empty:
            fig.add_annotation(x=str(te_df["CalYear"].iloc[0]),y=1.08,xref="x",yref="paper",
                               text="▼ Term End",showarrow=False,font=dict(color="#f39c12",size=10,family="Arial Bold"),
                               xanchor="center")
    fig.update_layout(barmode="stack",title=dict(text=title,y=0.97),
                      xaxis_title="Year",yaxis_title="($)",height=400,
                      legend=dict(orientation="h",yanchor="bottom",y=1.02,xanchor="left",x=0),
                      margin=dict(t=80,b=40))
    return fig

# ──────────────────────────────────────────────────────────────────
# WORD WIREFRAME EXPORT  (FIX #10)
# ──────────────────────────────────────────────────────────────────
def generate_wireframe_docx(base_info: dict) -> bytes:
    """Create a wireframe .docx showing app layout without any sensitive data."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        import subprocess, sys
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

    doc = Document()

    def set_cell_bg(cell, color_hex):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color_hex); tcPr.append(shd)

    def box(label, width=6.5, color="E8F0FE"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[ {label} ]")
        run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x33,0x33,0x88)
        run.font.bold = True
        from docx.oxml import OxmlElement as OE
        pPr = p._p.get_or_add_pPr()
        pBdr = OE("w:pBdr")
        for side in ["top","bottom","left","right"]:
            b = OE(f"w:{side}")
            b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4")
            b.set(qn("w:space"),"1"); b.set(qn("w:color"),"3355AA")
            pBdr.append(b)
        pPr.append(pBdr)
        p.paragraph_format.space_after = Pt(6)

    def section_title(t):
        p = doc.add_heading(t, level=2)
        p.runs[0].font.color.rgb = RGBColor(0x1a,0x3c,0x5e)

    def field_row(labels):
        tbl = doc.add_table(rows=1, cols=len(labels))
        tbl.style = "Table Grid"
        for i,lbl in enumerate(labels):
            cell = tbl.rows[0].cells[i]
            cell.text = lbl
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_bg(cell, "EEF2FF")
        doc.add_paragraph()

    def metric_row(labels):
        tbl = doc.add_table(rows=2, cols=len(labels))
        tbl.style = "Table Grid"
        for i,lbl in enumerate(labels):
            cell = tbl.rows[0].cells[i]
            cell.text = lbl
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            set_cell_bg(cell,"D5E8F0")
            val_cell = tbl.rows[1].cells[i]
            val_cell.text = "████"
            val_cell.paragraphs[0].runs[0].font.size = Pt(10)
            set_cell_bg(val_cell,"F5F8FF")
        doc.add_paragraph()

    # ── Title page ────────────────────────────────────────────────
    h = doc.add_heading("🏠 Canadian Mortgage Analyzer", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("App Wireframe — Structure reference (no sensitive data)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor(0x66,0x66,0x66)
    doc.add_paragraph(f"Generated: {date.today().strftime('%B %d, %Y')}")
    doc.add_page_break()

    tabs = [
        ("Tab 1: Setup & Overview", _wireframe_setup),
        ("Tab 2: Rate Change Scenarios", _wireframe_scenarios),
        ("Tab 3: Amortization Schedule", _wireframe_schedule),
        ("Tab 4: Prepayment Analysis", _wireframe_prepayment),
        ("Tab 5: Break Penalty", _wireframe_breakpen),
        ("Tab 6: Scenario Comparison", _wireframe_comparison),
    ]
    for tab_title, fn in tabs:
        doc.add_heading(tab_title, level=1)
        fn(doc, box, section_title, field_row, metric_row)
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def _wireframe_setup(doc,box,st_,fr,mr):
    st_("Section A: Purchase & Down Payment")
    fr(["Purchase Price ($)", "Down Payment (%)", "Down Payment ($)", "CMHC Premium"])
    box("CMHC advisory info bar (insured / not insured)")
    st_("Section B: Initial Mortgage Terms")
    fr(["Mortgage Type", "Payment Frequency", "Interest Rate (%)", "Amortization (yrs)"])
    fr(["Term (yrs)", "Mortgage Start Date", "", ""])
    st_("Section C: Additional Renewal Terms (collapsable)")
    box("➕ Add Past Renewal  |  [Renewal 1 row ▼]  |  [Renewal 2 row ▼]")
    st_("Section D: Past Prepayments (collapsable)")
    box("➕ Add Past Prepayment  |  [Prepayment 1 row ▼]")
    st_("Key Metrics at a Glance")
    mr(["Initial Principal","Balance @ Term End","Balance Today","Principal Paid"])
    mr(["Interest Paid","Remaining Amort","Total Interest","Original Amort"])
    mr(["Current Monthly Payment","","",""])
    st_("Charts")
    box("Donut Chart — Principal vs Total Interest", color="FFF0E0")
    box("Stacked Bar Chart — Yearly P & I (Past/Current/Post segments)", color="E0F0FF")
    box("💾 Save Setup to DB  button")

def _wireframe_scenarios(doc,box,st_,fr,mr):
    box("➕ New Scenario button")
    st_("Scenario expander (Scenario 1)")
    fr(["Name", "Description (50+ words text area)"])
    box("🚀 Quick templates checkbox → template selector + Apply button")
    box("➕ Add Renewal Entry button")
    st_("Renewal 1 row")
    fr(["Mode (By Date/Period)", "Effective Date", "Type (Fixed/Variable)", "Rate (%)","Term (yrs)"])
    box("Term start → end date caption")
    box("⚡ Early Renewal banner (if applicable)")
    fr(["Original Posted Rate", "Current Posted Rate"])
    box("Advisory penalty radio: 3-Month Interest | IRD | Custom value + text box")
    fr(["Misc Fees ($)"])
    box("Penalty/break-even summary info bar")
    st_("Variable Sub-scenarios (if Variable type)")
    box("➕ Add Sub-Scenario  |  [Sub-scenario a row]")
    st_("Scenario Results")
    mr(["Base Interest","Scenario Interest","Current Remaining","Scenario Remaining","Payment"])
    st_("Payment Override")
    fr(["Monthly Payment ($) — editable", "→ Adjusted Remaining", "→ Mortgage-free by"])
    box("Stacked Bar — P & I (Past/Current/Post), year x-axis, matching Image 2 style", color="E0F0FF")
    box("Rate over time chart (date x-axis)")
    box("💾 Save scenario button  |  Name input")
    doc.add_heading("Saved Scenarios Section (below new scenarios)", level=3)
    box("💾 Saved Scenarios header")
    box("Expander per scenario: 4 metrics | ✏️ Load for Editing | 🗑️ Delete")

def _wireframe_schedule(doc,box,st_,fr,mr):
    fr(["Scenario dropdown (Current Setup / Saved Scenarios)"])
    fr(["Hierarchy View toggle", "Highlight today checkbox"])
    st_("Hierarchy View Mode (default)")
    box("▶ PAST  [collapsed — click to expand]  | Total rows: N")
    box("▼ CURRENT (last 4 months + today)  [expanded]")
    box("    └─ Year 2025  └─ Nov 2025  └─ Dec 2025  └─ Jan 2026  └─ Feb 2026  └─ Mar 2026 (TODAY ★)")
    box("▶ FUTURE  [collapsed — click to expand]  | Remaining rows: N")
    st_("Full Table (when expanded)")
    box("Scrollable data table — header frozen (dark navy) — future rows in blue-italic")
    box("Current month row highlighted in yellow")
    st_("Balance Chart")
    box("Line chart: Balance + Cum Interest vs Date  |  'Today' vline", color="E0F0FF")
    box("⬇️ Download CSV button")

def _wireframe_prepayment(doc,box,st_,fr,mr):
    fr(["Rate scenario base dropdown"])
    st_("Left Column — Prepayment Options")
    fr(["Annual lump-sum ($)", "Month each year", "Starting year", "For how many years?"])
    fr(["Lender limit (%)", "", "", ""])
    st_("Right Column — Other Options")
    fr(["Increase type radio", "Extra per payment ($) / % increase"])
    fr(["One-time mode radio", "Date or Period #", "Amount ($)"])
    st_("Metrics")
    mr(["Interest (rate sc.)","Interest (+ prepayments)","Remaining (rate sc.)","Remaining (+ prepayments)","Total New Prepaid","Interest ROI"])
    box("Stacked Bar — P & I with prepayment impact", color="E0F0FF")
    box("💾 Save Prepayment Scenario  |  Name input")

def _wireframe_breakpen(doc,box,st_,fr,mr):
    st_("Inputs")
    fr(["Outstanding Balance ($)", "Contract Rate (%)", "Mortgage Type", "Months Remaining"])
    fr(["Misc Fees ($)", "", "", ""])
    fr(["Posted Rate at Origination (%)", "Current Posted Rate for Remaining Term (%)"])
    st_("Advisory Calculation")
    box("Radio: 3-Month Interest ($X) | IRD ($X) | Custom value [____]")
    mr(["3 Months Interest","IRD","Penalty Applied","Total Exit Cost"])
    st_("Break-even Analysis")
    fr(["New rate if you break (%) slider"])
    mr(["Interest (Stay)","Interest+Fees (Break)","Net Savings"])
    box("Line chart: Net Savings vs New Rate  |  break-even line  |  current rate marker", color="E0F0FF")
    box("Monthly payment change + months to recoup info bar")

def _wireframe_comparison(doc,box,st_,fr,mr):
    box("Scenarios selector: 2 / 3 / 4 radio")
    st_("Scenario columns (e.g. 2 columns)")
    fr(["Label", "Rate (%)", "Amort (yrs)", "Frequency", "Annual Lump ($)"])
    doc.add_paragraph("^ Scenario 1 column  |  ^ Scenario 2 column (repeated per scenario)")
    doc.add_paragraph()
    box("Note: First scenario defaults to 'Current Mortgage'  — can be changed")
    st_("Results")
    mr(["Scenario","Rate","Amort","Freq","Orig Payment","Current Payment","Remaining","Total Interest","Total Paid"])
    box("Balance comparison chart — date x-axis, one line per scenario", color="E0F0FF")
    box("Stacked bar P & I by scenario and year", color="E0F0FF")
    box("🏆 Best scenario callout  |  📥 Download prompt (.txt)")

# ──────────────────────────────────────────────────────────────────
# SESSION STATE
# ──────────────────────────────────────────────────────────────────
for k,v in {"db_conn":None,"setup_loaded":False,"setup_data":None,
            "rc_scenarios":{},"past_prepayments":[],"past_renewals":[]}.items():
    if k not in st.session_state: st.session_state[k]=v

# ══════════════════════════════════════════════════════════════════
# GATE: DB CONNECTION
# ══════════════════════════════════════════════════════════════════
if not st.session_state.db_conn:
    st.markdown('<div class="main-header">🏠 Canadian Mortgage Analyzer</div>',unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Connect to your SQL Server database to begin.</div>',unsafe_allow_html=True)
    _,cm,_=st.columns([1,2,1])
    with cm:
        st.markdown('<div class="db-gate">',unsafe_allow_html=True)
        st.markdown("### 🗄️ Database Connection")
        srv=st.text_input("SQL Server",r"localhost\SQLEXPRESS",key="g_srv")
        db =st.text_input("Database","MortgageDB",key="g_db")
        tru=st.checkbox("Windows Authentication",True,key="g_tru")
        usr=pwd=""
        if not tru:
            usr=st.text_input("Username",key="g_usr")
            pwd=st.text_input("Password",type="password",key="g_pwd")
        if st.button("🔌 Connect & Continue",use_container_width=True,key="btn_gate"):
            conn,err=get_db_connection(srv,db,tru,usr,pwd)
            if conn:
                st.session_state.db_conn=conn
                ex=db_load_setup(conn)
                if ex:
                    st.session_state.setup_data=ex; st.session_state.setup_loaded=True
                    st.session_state.past_renewals=ex.get("past_renewals",[])
                    st.session_state.past_prepayments=ex.get("past_prepayments",[])
                st.rerun()
            else: st.error(f"❌ {err}")
        st.markdown("</div>",unsafe_allow_html=True)
    st.stop()

if not st.session_state.setup_loaded:
    ex=db_load_setup(st.session_state.db_conn)
    if ex:
        st.session_state.setup_data=ex; st.session_state.setup_loaded=True
        st.session_state.past_renewals=ex.get("past_renewals",[])
        st.session_state.past_prepayments=ex.get("past_prepayments",[])

# ── Header ────────────────────────────────────────────────────────
hc1,hc2=st.columns([5,1])
hc1.markdown('<div class="main-header">🏠 Canadian Mortgage Analyzer</div>',unsafe_allow_html=True)
hc1.markdown('<div class="sub-header">Canadian semi-annual compounding · CMHC · Prepayments · Rate scenarios · Break penalties</div>',unsafe_allow_html=True)
hc2.markdown("<div style='text-align:right;padding-top:1.2rem;color:#27ae60;font-size:.85rem;'>🟢 DB Connected</div>",unsafe_allow_html=True)

def require_setup():
    if not st.session_state.get("base"):
        st.info("⬅️ Complete **Setup & Overview** tab and click 💾 Save Setup to DB first.")
        st.stop()

# FIX #6: Tab order — Setup | Rate Change | Amortization | Prepayment | Break Penalty | Comparison
tabs=st.tabs(["📊 Setup & Overview","📈 Rate Change Scenarios",
              "📅 Amortization Schedule","💰 Prepayment Analysis",
              "⚠️ Break Penalty","🔄 Scenario Comparison"])

# ══════════════════════════════════════════════════════════════════
# TAB 1 — SETUP & OVERVIEW  (FIX #2: split into sections)
# ══════════════════════════════════════════════════════════════════
with tabs[0]:
    st.subheader("Mortgage Setup")
    sd=st.session_state.setup_data or {}
    def sv(key,default):
        return sd.get("widget_state",{}).get(key,default)

    # ── SECTION A: Purchase & Down Payment ───────────────────────
    st.markdown("#### 🏡 A · Purchase & Down Payment")
    a1,a2=st.columns(2)
    purchase_price=a1.number_input("Purchase Price ($)",100_000,5_000_000,int(sv("s_price",1_030_000)),5_000,format="%d",key="s_price",help="Total property purchase price")
    down_pct=a2.slider("Down Payment (%)",5.0,50.0,float(sv("s_dpct",20.0)),0.5,key="s_dpct",help="Down payment as % of purchase price")
    down_pay=purchase_price*down_pct/100; a1.metric("Down Payment",f"${down_pay:,.0f}")

    cmhc,hst=cmhc_premium(purchase_price,down_pay)
    if cmhc is None:
        st.markdown('<div class="warn">⚠️ CMHC not available (price >$1.5M or down <5%)</div>',unsafe_allow_html=True)
        insured_p=purchase_price-down_pay
    elif cmhc==0:
        st.markdown('<div class="ok">✅ No CMHC premium — down ≥ 20%</div>',unsafe_allow_html=True)
        insured_p=purchase_price-down_pay
    else:
        add_c=a2.checkbox("Add CMHC to mortgage?",bool(sv("s_addcmhc",True)),key="s_addcmhc")
        st.markdown(f'<div class="inf">🛡️ CMHC: <b>${cmhc:,.0f}</b> (+HST ~${hst:,.0f}) · {cmhc/(purchase_price-down_pay)*100:.2f}%</div>',unsafe_allow_html=True)
        insured_p=(purchase_price-down_pay)+(cmhc if add_c else 0)

    st.divider()

    # ── SECTION B: Initial Mortgage Terms ────────────────────────
    st.markdown("#### 💵 B · Initial Mortgage Terms")
    b1,b2=st.columns(2)
    mortgage_type=b1.selectbox("Mortgage Type",["Fixed","Variable"],index=["Fixed","Variable"].index(sv("s_mtype","Fixed")),key="s_mtype",help="Fixed = rate locked for term; Variable = moves with prime")
    payment_freq=b2.selectbox("Payment Frequency",list(FREQ.keys()),index=list(FREQ.keys()).index(sv("s_freq","Monthly")),key="s_freq")
    annual_rate=b1.number_input("Interest Rate (%)",0.5,20.0,float(sv("s_rate",5.39)),0.01,format="%.2f",key="s_rate",help="Your contracted annual interest rate")
    amort_years=b2.slider("Amortization (years)",5,30,int(sv("s_amort",30)),key="s_amort",help="Total years to pay off the mortgage from origination")
    term_opts=[0.5,1,2,3,4,5,7,10]
    term_years=b1.selectbox("Term (years)",term_opts,index=term_opts.index(sv("s_term",3)) if sv("s_term",3) in term_opts else 3,key="s_term",help="Rate-locked period length")
    _sd_raw=sv("s_startdate","2023-08-15")
    _sd_val=date.fromisoformat(_sd_raw) if isinstance(_sd_raw,str) else _sd_raw
    start_date_in=b2.date_input("Mortgage Start Date",_sd_val,key="s_startdate",help="Date your mortgage originally started")
    if down_pct<20 and amort_years>25:
        st.markdown('<div class="warn">⚠️ Insured mortgages limited to 25-yr amortization.</div>',unsafe_allow_html=True)

    fc=FREQ[payment_freq]; n_py=fc["n"]; accel=fc["accel"]

    # ── SECTION C: Additional Renewal Terms (collapsable) ─────────
    st.divider()
    with st.expander("🔄 C · Additional Renewal Terms (Past Renewals)",expanded=False):
        st.caption("Add renewal terms that have already taken effect since your mortgage started.")
        if st.button("➕ Add Past Renewal",key="btn_add_rn"):
            if st.session_state.past_renewals:
                last=st.session_state.past_renewals[-1]
                prev_end=date.fromisoformat(last["start_date_str"])+relativedelta(years=int(last["term_years"]),months=int((float(last["term_years"])%1)*12))
            else:
                prev_end=start_date_in+relativedelta(years=int(term_years),months=int((term_years%1)*12))
            st.session_state.past_renewals.append({"id":str(uuid.uuid4())[:8],"start_date_str":str(prev_end),"rate":annual_rate,"mtype":"Fixed","term_years":3})
            st.rerun()
        del_rn=[]
        for idx,rn in enumerate(st.session_state.past_renewals):
            rr=st.columns([2,1.5,1.5,1.5,0.7])
            nsd=rr[0].date_input(f"Start #{idx+1}",date.fromisoformat(rn["start_date_str"]),key=f"rn_sd_{rn['id']}")
            nr =rr[1].number_input(f"Rate #{idx+1} (%)",0.5,20.0,float(rn["rate"]),0.01,format="%.2f",key=f"rn_rt_{rn['id']}")
            nmt=rr[2].selectbox(f"Type #{idx+1}",["Fixed","Variable"],index=0 if rn["mtype"]=="Fixed" else 1,key=f"rn_mt_{rn['id']}")
            nty=rr[3].selectbox(f"Term #{idx+1}",term_opts,index=term_opts.index(rn["term_years"]) if rn["term_years"] in term_opts else 3,key=f"rn_ty_{rn['id']}")
            if rr[4].button("🗑️",key=f"del_rn_{rn['id']}"): del_rn.append(idx)
            end_d=date.fromisoformat(str(nsd))+relativedelta(years=int(nty),months=int((float(nty)%1)*12))
            rr[0].caption(f"End: {end_d.strftime('%b %Y')}")
            st.session_state.past_renewals[idx].update(start_date_str=str(nsd),rate=float(nr),mtype=nmt,term_years=nty)
        for i in sorted(del_rn,reverse=True): st.session_state.past_renewals.pop(i)
        if del_rn: st.rerun()

    # ── SECTION D: Past Prepayments (collapsable) ─────────────────
    with st.expander("💳 D · Past Prepayments Already Made",expanded=False):
        if st.button("➕ Add Past Prepayment",key="btn_add_pp"):
            st.session_state.past_prepayments.append({"id":str(uuid.uuid4())[:8],"date_str":str(start_date_in),"amount":0.0})
            st.rerun()
        del_pp=[]
        for idx,pp in enumerate(st.session_state.past_prepayments):
            r=st.columns([2,2,1])
            nd=r[0].date_input(f"Date #{idx+1}",date.fromisoformat(pp["date_str"]),min_value=start_date_in,max_value=date.today(),key=f"ppd_{pp['id']}")
            na=r[1].number_input(f"Amount ($) #{idx+1}",0,2_000_000,int(pp["amount"]),500,key=f"ppa_{pp['id']}")
            if r[2].button("🗑️",key=f"del_pp_{pp['id']}"): del_pp.append(idx)
            st.session_state.past_prepayments[idx].update(date_str=str(nd),amount=float(na))
        for i in sorted(del_pp,reverse=True): st.session_state.past_prepayments.pop(i)
        if del_pp: st.rerun()

    # ── Build schedules ───────────────────────────────────────────
    past_renewal_rcs=[{"period":date_to_period(rn["start_date_str"],start_date_in,n_py),"new_rate":float(rn["rate"])} for rn in st.session_state.past_renewals]
    past_extra=[{"period":date_to_period(pp["date_str"],start_date_in,n_py),"amount":float(pp["amount"])} for pp in st.session_state.past_prepayments if pp["amount"]>0]
    full_df,full_sum=build_amortization(insured_p,annual_rate,n_py,amort_years,accel=accel,start_date=start_date_in,extra_payments=past_extra or None,rate_changes=past_renewal_rcs or None)
    today_m=get_today_metrics(full_df,n_py)
    # Current term end period
    orig_term_end_p=int(term_years*n_py)
    _,t_sum=build_amortization(insured_p,annual_rate,n_py,amort_years,accel=accel,start_date=start_date_in,extra_payments=past_extra or None,rate_changes=past_renewal_rcs or None,term_periods=orig_term_end_p)
    term_end_d=start_date_in+relativedelta(years=int(term_years),months=int((term_years%1)*12))

    # ── KEY METRICS  (FIX #3, #4) ─────────────────────────────────
    st.divider()
    st.markdown("#### 📊 Key Metrics at a Glance")

    rem_y=today_m.get("remaining_years",0)
    rem_end=today_m.get("remaining_end_date","")
    balance_today=today_m.get("balance_today",insured_p)
    # Current payment = payment on REMAINING balance at CURRENT rate for remaining amort
    current_rate_now=float(past_renewal_rcs[-1]["new_rate"]) if past_renewal_rcs else annual_rate
    curr_pmt=calc_pmt(balance_today,current_rate_now,n_py,rem_y,accel) if rem_y>0 else calc_pmt(insured_p,annual_rate,n_py,amort_years,accel)

    mc1,mc2,mc3=st.columns(3)
    mc1.markdown(f"""
    <div class="mc"><h3>Initial Mortgage Principal</h3><p>${insured_p:,.0f}</p></div>
    <div class="mc mc-b"><h3>Expected Balance at Term End ({term_end_d.strftime('%b %Y')})</h3><p>${t_sum.get('end_balance',insured_p):,.0f}</p></div>
    <div class="mc"><h3>Original Amortization Period</h3><p>{amort_years} years</p></div>
    """,unsafe_allow_html=True)

    mc2.markdown(f"""
    <div class="mc mc-g"><h3>🟢 Balance as of Today ({today_m.get('as_of_date','')})</h3><p>${balance_today:,.0f}</p></div>
    <div class="mc mc-g"><h3>🟢 Principal Paid to Date</h3><p>${today_m.get('principal_paid_today',0):,.0f}</p></div>
    <div class="mc mc-g"><h3>🟢 Interest Paid to Date</h3><p>${today_m.get('interest_paid_today',0):,.0f}</p></div>
    """,unsafe_allow_html=True)

    mc3.markdown(f"""
    <div class="mc"><h3>⏳ Current Remaining Amortization</h3>
        <p>{amort_years} yrs original &nbsp;·&nbsp; {rem_y:.1f} yrs more &nbsp;·&nbsp; ends {rem_end}</p></div>
    <div class="mc mc-r"><h3>Total Interest (full remaining amortization)</h3><p>${full_sum.get('total_interest',0):,.0f}</p></div>
    <div class="mc"><h3>📆 Current Monthly Payment</h3><p>${curr_pmt:,.2f}</p></div>
    """,unsafe_allow_html=True)

    # ── Charts ────────────────────────────────────────────────────
    st.divider()
    cc1,cc2=st.columns(2)
    with cc1:
        fig_d=go.Figure(go.Pie(labels=["Principal","Total Interest"],values=[insured_p,full_sum.get("total_interest",0)],hole=0.55,marker_colors=["#1a3c5e","#e74c3c"],textinfo="label+percent"))
        fig_d.update_layout(title="Principal vs Total Interest",height=280,margin=dict(t=40,b=5))
        st.plotly_chart(fig_d,use_container_width=True,key="ch_donut")
    with cc2:
        today_p_g=today_m.get("period_today",0)
        fig_pi=stacked_bar_pi(full_df,today_p_g,orig_term_end_p,"Yearly P & I — 3 Segments")
        st.plotly_chart(fig_pi,use_container_width=True,key="ch_pi")

    # ── Save Setup + Wire export ───────────────────────────────────
    st.divider()
    sv1,sv2,sv3=st.columns([1,1,2])
    if sv1.button("💾 Save Setup to DB",key="btn_ss"):
        payload={"widget_state":{"s_price":purchase_price,"s_dpct":down_pct,"s_mtype":mortgage_type,"s_freq":payment_freq,"s_rate":annual_rate,"s_amort":amort_years,"s_term":term_years,"s_startdate":str(start_date_in),"s_addcmhc":True},"past_renewals":st.session_state.past_renewals,"past_prepayments":st.session_state.past_prepayments,"summary":full_sum,"today_metrics":today_m}
        if db_save_setup(st.session_state.db_conn,payload):
            st.session_state.setup_data=payload; st.session_state.setup_loaded=True
            sv2.success("✅ Saved to database.")
        else: sv2.error("❌ Failed to save.")

    # FIX #10: Word wireframe export
    if sv2.button("📄 Export App Wireframe (.docx)",key="btn_wire"):
        with st.spinner("Generating wireframe (installing python-docx if needed)..."):
            try:
                wire_bytes=generate_wireframe_docx({})
                st.session_state["wire_bytes"]=wire_bytes
            except Exception as ex_:
                st.error(f"❌ Wireframe generation failed: {ex_}. Try: pip install python-docx")
    if st.session_state.get("wire_bytes"):
        sv3.download_button("⬇️ Download Wireframe",data=st.session_state["wire_bytes"],
            file_name="mortgage_analyzer_wireframe.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="btn_wire_dl")

    st.session_state["base"]=dict(principal=insured_p,annual_rate=annual_rate,n_py=n_py,amort_years=amort_years,accel=accel,start_date=start_date_in,mortgage_type=mortgage_type,term_years=term_years,payment_freq=payment_freq,purchase_price=purchase_price,down_payment=down_pay,past_extra=past_extra,past_renewal_rcs=past_renewal_rcs,today_m=today_m,orig_term_end_p=orig_term_end_p,current_rate=current_rate_now,curr_pmt=curr_pmt,term_end_d=term_end_d,full_sum=full_sum)

# ══════════════════════════════════════════════════════════════════
# TAB 2 — RATE CHANGE SCENARIOS  (FIX #6: now second tab)
# ══════════════════════════════════════════════════════════════════
with tabs[1]:
    st.subheader("📈 Rate Change / Renewal Scenarios")
    require_setup()
    b=st.session_state["base"]

    st.info("Create named renewal scenarios. Early renewals auto-calculate break penalty (advisory). Variable renewals support sub-scenarios a/b/c.")

    # FIX #8: Only DB scenarios — load fresh each render
    db_scenarios=db_load_scenarios(st.session_state.db_conn)
    # Map by name for quick lookup
    db_sc_by_name={s["name"]:s for s in db_scenarios}

    rcs:dict=st.session_state.rc_scenarios
    if st.button("➕ New Scenario",key="btn_new_rc"):
        nid=str(uuid.uuid4())[:8]
        rcs[nid]={"name":f"Scenario {len(rcs)+1}","desc":"","renewals":[]}
        st.rerun()
    if not rcs:
        st.markdown('<div class="inf">Click ➕ New Scenario to begin.</div>',unsafe_allow_html=True)

    orig_term_end_p_sc=b["orig_term_end_p"]
    df_base_ref,s_base_ref=build_amortization(b["principal"],b["annual_rate"],b["n_py"],b["amort_years"],accel=b["accel"],start_date=b["start_date"],extra_payments=b.get("past_extra") or None,rate_changes=b.get("past_renewal_rcs") or None)
    today_p_ref=b["today_m"].get("period_today",0)
    # FIX #7: base remaining = current remaining amortization (not original amort_years)
    base_remaining_yrs=b["today_m"].get("remaining_years",b["amort_years"])
    term_opts_sc=[0.5,1,2,3,4,5,7,10]
    sc_del=[]

    for sc_id,sc in rcs.items():
        with st.expander(f"📋 {sc['name']}"+(f" — {sc['desc'][:60]}" if sc["desc"] else ""),expanded=True):
            h1,h2,h3=st.columns([2,3,1])
            sc["name"]=h1.text_input("Name",sc["name"],key=f"rcn_{sc_id}")
            sc["desc"]=h2.text_area("Description",sc["desc"],height=80,placeholder="Describe this scenario in detail — e.g. Early renewal at a lower rate before the current term ends, anticipating Bank of Canada cuts by mid-2025, followed by a gradual rise through the next 3-year fixed term ending in 2028.",key=f"rcd_{sc_id}")
            if h3.button("🗑️ Delete",key=f"del_sc_{sc_id}"): sc_del.append(sc_id)

            show_tpl=st.checkbox("🚀 Quick templates",False,key=f"tpl_cb_{sc_id}")
            if show_tpl:
                ren_p=orig_term_end_p_sc+1
                ren_d=str(period_to_date(ren_p,b["start_date"],b["n_py"]))
                tpls={"+1% at renewal":[{"date":ren_d,"rate":b["annual_rate"]+1,"mtype":"Fixed","term":3}],"+2% at renewal":[{"date":ren_d,"rate":b["annual_rate"]+2,"mtype":"Fixed","term":3}],"-1% at renewal":[{"date":ren_d,"rate":b["annual_rate"]-1,"mtype":"Fixed","term":3}],"-2% at renewal":[{"date":ren_d,"rate":b["annual_rate"]-2,"mtype":"Fixed","term":3}],"Variable at renewal":[{"date":ren_d,"rate":b["annual_rate"]-0.5,"mtype":"Variable","term":3}],"BoC hike then cut":[{"date":str(period_to_date(ren_p//2,b["start_date"],b["n_py"])),"rate":b["annual_rate"]+2,"mtype":"Fixed","term":1},{"date":ren_d,"rate":b["annual_rate"]+1,"mtype":"Fixed","term":3}],"Rate stays flat":[{"date":ren_d,"rate":b["annual_rate"],"mtype":"Fixed","term":3}]}
                tc1,tc2=st.columns([3,1])
                tpl_s=tc1.selectbox("Template",list(tpls.keys()),key=f"tpl_sel_{sc_id}")
                if tc2.button("Apply",key=f"tpl_ap_{sc_id}"):
                    sc["renewals"]=[{"id":str(uuid.uuid4())[:8],"mode":"By Date","date_str":t["date"],"period":date_to_period(t["date"],b["start_date"],b["n_py"]),"new_rate":t["rate"],"mtype":t["mtype"],"term_years":t["term"],"actual_penalty":0,"misc_fees":250,"orig_posted":t["rate"]+1.5,"curr_posted":t["rate"]-0.5,"variable_subs":{}} for t in tpls[tpl_s]]
                    st.rerun()

            st.markdown("---")
            if st.button("➕ Add Renewal Entry",key=f"add_ren_{sc_id}"):
                dd=str(start_date_in+relativedelta(years=int(b["term_years"]),months=int((b["term_years"]%1)*12)))
                sc["renewals"].append({"id":str(uuid.uuid4())[:8],"mode":"By Date","date_str":dd,"period":date_to_period(dd,b["start_date"],b["n_py"]),"new_rate":b["annual_rate"],"mtype":"Fixed","term_years":3,"actual_penalty":0,"misc_fees":250,"orig_posted":b["annual_rate"]+1.5,"curr_posted":b["annual_rate"]-0.5,"variable_subs":{}})
                st.rerun()

            prev_term_end_p=orig_term_end_p_sc
            ren_del=[]

            for ri,rn in enumerate(sc["renewals"]):
                rid=rn["id"]
                st.markdown(f"**Renewal {ri+1}**")
                rc1,rc2,rc3,rc4,rc5=st.columns([1.5,1.8,1.5,1.5,0.7])
                rn["mode"]=rc1.radio("Mode",["By Date","By Period"],index=0 if rn.get("mode","By Date")=="By Date" else 1,horizontal=True,key=f"rm_{sc_id}_{rid}")
                if rn["mode"]=="By Date":
                    pd_v=rc2.date_input("Effective date",date.fromisoformat(rn.get("date_str",str(b["start_date"]))),key=f"rd_{sc_id}_{rid}")
                    rn["date_str"]=str(pd_v); rn["period"]=date_to_period(pd_v,b["start_date"],b["n_py"])
                    rc2.caption(f"≈ Period {rn['period']}")
                else:
                    mx=int(b["amort_years"]*b["n_py"])
                    rn["period"]=int(rc2.number_input("Period #",1,mx,int(rn.get("period",orig_term_end_p_sc+1)),key=f"rp_{sc_id}_{rid}"))
                    rc2.caption(f"≈ {period_to_date(rn['period'],b['start_date'],b['n_py']).strftime('%b %Y')}")

                rn["mtype"]=rc3.selectbox("Type",["Fixed","Variable"],index=0 if rn.get("mtype","Fixed")=="Fixed" else 1,key=f"rmt_{sc_id}_{rid}")
                rn["new_rate"]=float(rc4.number_input("Rate (%)",0.5,20.0,float(rn.get("new_rate",b["annual_rate"])),0.01,format="%.2f",key=f"rrt_{sc_id}_{rid}",help="New interest rate at this renewal"))
                if rc5.button("🗑️",key=f"delren_{sc_id}_{rid}"): ren_del.append(ri)

                rn["term_years"]=st.selectbox(f"Term (years) — Renewal {ri+1}",term_opts_sc,index=term_opts_sc.index(rn.get("term_years",3)) if rn.get("term_years",3) in term_opts_sc else 3,key=f"rty_{sc_id}_{rid}")
                rn_start_d=period_to_date(rn["period"],b["start_date"],b["n_py"])
                rn_end_d=rn_start_d+relativedelta(years=int(rn["term_years"]),months=int((float(rn["term_years"])%1)*12))
                st.caption(f"📅 Term: **{rn_start_d.strftime('%b %d, %Y')}** → **{rn_end_d.strftime('%b %d, %Y')}**")

                is_early=rn["period"]<prev_term_end_p
                months_left_at=max(int((prev_term_end_p-rn["period"])/b["n_py"]*12),1) if is_early else 0
                if is_early:
                    ren_df=df_base_ref[df_base_ref["Period"]<=rn["period"]]
                    bal_ren=float(ren_df["Balance"].iloc[-1]) if not ren_df.empty else b["principal"]
                    rate_ren=float(ren_df["Rate (%)"].iloc[-1]) if not ren_df.empty else b["annual_rate"]
                    st.markdown(f'<div class="warn">⚡ <b>Early Renewal</b> — {months_left_at} months remain · Balance: <b>${bal_ren:,.0f}</b></div>',unsafe_allow_html=True)
                    bp1,bp2=st.columns(2)
                    rn["orig_posted"]=float(bp1.number_input("Original posted rate (%)",0.5,20.0,float(rn.get("orig_posted",rate_ren+1.5)),0.01,format="%.2f",key=f"op_{sc_id}_{rid}",help="Bank's posted rate at origination"))
                    rn["curr_posted"]=float(bp2.number_input("Current posted rate (%)",0.5,20.0,float(rn.get("curr_posted",max(rate_ren-0.5,0.5))),0.01,format="%.2f",key=f"cp_{sc_id}_{rid}",help="Current posted rate for remaining term"))
                    adv=calc_break_penalty(bal_ren,rate_ren,rn["mtype"],rn["orig_posted"],rn["curr_posted"],months_left_at)
                    pen_opts=[f"3-Month Interest (${adv['3_months_interest']:,.0f})"]
                    if adv["ird"] is not None: pen_opts.append(f"IRD (${adv['ird']:,.0f})")
                    pen_opts.append("Custom value")
                    pc1,pc2=st.columns([3,2])
                    pen_choice=pc1.radio("Apply which penalty?",pen_opts,key=f"pen_radio_{sc_id}_{rid}",help="Advisory — your bank may charge differently")
                    if "Custom" in pen_choice:
                        custom_str=pc2.text_input("",value=str(int(adv["calc_penalty"])),key=f"cpen_{sc_id}_{rid}")
                        pc2.caption("Custom value ($)")
                        try: rn["actual_penalty"]=float(custom_str.replace(",","").replace("$",""))
                        except: rn["actual_penalty"]=adv["calc_penalty"]
                    elif "IRD" in pen_choice: rn["actual_penalty"]=adv["ird"] or 0.0
                    else: rn["actual_penalty"]=adv["3_months_interest"]
                    rn["misc_fees"]=float(st.number_input("Misc fees ($)",0,50_000,int(rn.get("misc_fees",500)),50,key=f"mf_{sc_id}_{rid}"))
                    total_exit=rn["actual_penalty"]+rn["misc_fees"]
                    old_pmt=calc_pmt(bal_ren,rate_ren,12,max(b["amort_years"]-rn["period"]/b["n_py"],1))
                    new_pmt=calc_pmt(bal_ren,rn["new_rate"],12,max(b["amort_years"]-rn["period"]/b["n_py"],1))
                    st.markdown(f'<div class="pen">💸 Penalty applied: <b>${rn["actual_penalty"]:,.0f}</b> · Misc: <b>${rn["misc_fees"]:,.0f}</b> · <b>Total: ${total_exit:,.0f}</b></div>',unsafe_allow_html=True)
                    if abs(old_pmt-new_pmt)>1:
                        rec=total_exit/abs(old_pmt-new_pmt)
                        st.caption(f"Monthly saving: ${abs(old_pmt-new_pmt):,.0f} → Break-even: **{rec:.0f} months** ({rec/12:.1f} yrs)")
                else:
                    rn["misc_fees"]=float(st.number_input("Misc fees ($)",0,50_000,int(rn.get("misc_fees",250)),50,key=f"mf2_{sc_id}_{rid}"))
                    rn["actual_penalty"]=0

                if rn["mtype"]=="Variable":
                    if "variable_subs" not in rn: rn["variable_subs"]={}
                    vsubs=rn["variable_subs"]
                    if st.button("➕ Add Sub-Scenario",key=f"add_vsub_{sc_id}_{rid}"):
                        letter=chr(ord('a')+len(vsubs))
                        vsubs[letter]={"name":f"Sub {letter}","n_changes":1,"changes":[]}
                        st.rerun()
                    vsub_del=[]
                    for sub_k,sub in vsubs.items():
                        vsa,vsb,vsc_=st.columns([2,1,1])
                        sub["name"]=vsa.text_input("Name",sub["name"],key=f"vsn_{sc_id}_{rid}_{sub_k}")
                        sub["n_changes"]=int(vsb.number_input("# changes",1,12,int(sub.get("n_changes",1)),1,key=f"vsnc_{sc_id}_{rid}_{sub_k}"))
                        if vsc_.button("🗑️",key=f"del_vsub_{sc_id}_{rid}_{sub_k}"): vsub_del.append(sub_k)
                        while len(sub["changes"])<sub["n_changes"]: sub["changes"].append({"id":str(uuid.uuid4())[:8],"date_str":rn["date_str"],"new_rate":rn["new_rate"]})
                        sub["changes"]=sub["changes"][:sub["n_changes"]]
                        for ci_,chg in enumerate(sub["changes"]):
                            vc1,vc2=st.columns(2)
                            chg_d=vc1.date_input(f"Change {ci_+1} — date",value=date.fromisoformat(chg.get("date_str",rn["date_str"])),key=f"vcd_{sc_id}_{rid}_{sub_k}_{ci_}")
                            chg["date_str"]=str(chg_d); chg["period"]=date_to_period(chg_d,b["start_date"],b["n_py"])
                            chg["new_rate"]=float(vc2.number_input(f"Rate (%) — {ci_+1}",0.5,20.0,float(chg.get("new_rate",rn["new_rate"])),0.01,format="%.2f",key=f"vcr_{sc_id}_{rid}_{sub_k}_{ci_}"))
                    for sk in vsub_del: del vsubs[sk]
                    if vsub_del: st.rerun()

                prev_term_end_p=int(rn["period"])+int(float(rn["term_years"])*b["n_py"])
                st.markdown("---")

            for ri in sorted(ren_del,reverse=True): sc["renewals"].pop(ri)
            if ren_del: st.rerun()

            if sc["renewals"]:
                last_rn=sc["renewals"][-1]
                sc_term_end_p=int(last_rn["period"])+int(float(last_rn["term_years"])*b["n_py"])
            else: sc_term_end_p=orig_term_end_p_sc

            main_rc=[{"period":rn["period"],"new_rate":rn["new_rate"]} for rn in sc["renewals"]]
            all_rcs_sc=(b.get("past_renewal_rcs") or [])+main_rc
            df_sc,s_sc=build_amortization(b["principal"],b["annual_rate"],b["n_py"],b["amort_years"],accel=b["accel"],start_date=b["start_date"],extra_payments=b.get("past_extra") or None,rate_changes=all_rcs_sc or None)

            tm=b["today_m"]; today_p_sc=tm.get("period_today",0)
            rem_sc=round((len(df_sc)-today_p_sc)/b["n_py"],1) if today_p_sc>0 and not df_sc.empty else b["amort_years"]

            last_rate=sc["renewals"][-1]["new_rate"] if sc["renewals"] else b["annual_rate"]
            today_bal=tm.get("balance_today",b["principal"])
            calc_monthly_sc=calc_pmt(today_bal,last_rate,b["n_py"],rem_sc,b["accel"]) if rem_sc>0 else 0

            # FIX 2+3+5: metrics show adjusted values from payment input
            # Render payment input FIRST so user_pmt is available for metrics
            st.markdown("##### 💳 Payment & Amortization Impact")
            min_pmt=float(max(today_bal*periodic_rate(last_rate,b["n_py"])+1,100))
            pay1,pay2,pay3=st.columns([2,2,2])
            user_pmt=pay1.number_input(
                "Monthly payment ($)",
                min_value=min_pmt, max_value=float(today_bal),
                value=round(calc_monthly_sc,2), step=50.0, format="%.2f",
                key=f"user_pmt_{sc_id}",
                help=f"Calculated payment to maintain current amortization: ${calc_monthly_sc:,.2f}. Increase to pay off faster, decrease to extend.")
            pay1.caption(f"Payment to keep current amortization: **${calc_monthly_sc:,.2f}**")

            # Compute adjusted remaining from user_pmt
            adj_rem = calc_remaining_years(today_bal, last_rate, b["n_py"], user_pmt) if user_pmt > 0 and today_bal > 0 else rem_sc
            adj_end = date.today() + relativedelta(years=int(adj_rem), months=int((adj_rem%1)*12))
            pmt_changed = abs(user_pmt - calc_monthly_sc) > 0.5

            if pmt_changed:
                colour = "mc-g" if user_pmt > calc_monthly_sc else "mc-r"
                pay2.markdown(f'<div class="mc {colour}"><h3>Adjusted Remaining</h3><p>{adj_rem:.1f} yrs</p></div>', unsafe_allow_html=True)
                pay3.markdown(f'<div class="mc {colour}"><h3>Mortgage-free by</h3><p>{adj_end.strftime("%b %Y")}</p></div>', unsafe_allow_html=True)

            # FIX 3: Scenario Remaining = base_remaining when payment is the calculated payment
            display_rem_sc = adj_rem  # uses user_pmt; if unchanged == rem_sc == base_remaining_yrs

            m1,m2,m3,m4,m5=st.columns(5)
            m1.metric("Base Interest",f"${s_base_ref.get('total_interest',0):,.0f}",help="Total interest with no rate changes at current rate")
            # FIX 2: Scenario Interest delta shown alongside mortgage-free delta
            sc_int_delta=s_sc.get('total_interest',0)-s_base_ref.get('total_interest',0)
            m2.metric("Scenario Interest",f"${s_sc.get('total_interest',0):,.0f}",
                      delta=f"${sc_int_delta:+,.0f}",
                      help="Total interest under this renewal scenario")
            m3.metric("Current Remaining",f"{base_remaining_yrs:.1f} yrs",
                      help="Current remaining amortization from today at base rate")
            # FIX 2+3: Scenario Remaining driven by payment input
            sc_rem_delta = round(display_rem_sc - base_remaining_yrs, 1)
            m4.metric("Scenario Remaining",f"{display_rem_sc:.1f} yrs",
                      delta=f"{sc_rem_delta:+.1f} yrs" if sc_rem_delta != 0 else "Same as current",
                      help="Remaining amortization at this scenario's rate with chosen payment")
            # FIX 5: label changed
            m5.metric("Required Monthly Payment",f"${calc_monthly_sc:,.2f}",
                      help="Monthly payment required to maintain the current remaining amortization period at this scenario's interest rate")

            # FIX #1: Combined stacked bar (Image 2 style)
            if not df_sc.empty:
                fig_sc_bar=stacked_bar_pi(df_sc,today_p_sc,sc_term_end_p,f"{sc['name']} — P & I Breakdown")
                st.plotly_chart(fig_sc_bar,use_container_width=True,key=f"ch_sc_bar_{sc_id}")
                fig_rr=go.Figure()
                fig_rr.add_scatter(x=df_sc["Date"],y=df_sc["Rate (%)"],fill="tozeroy",name="Rate",line=dict(color="#27ae60"))
                fig_rr.update_layout(title="Rate over time",xaxis_title="Date",yaxis_title="%",height=200,margin=dict(t=40,b=30))
                st.plotly_chart(fig_rr,use_container_width=True,key=f"ch_rate_{sc_id}")

            # Variable sub-scenario charts
            for rn in sc["renewals"]:
                if rn["mtype"]!="Variable": continue
                for sub_k,sub in rn.get("variable_subs",{}).items():
                    sub_rc=all_rcs_sc+[{"period":chg["period"],"new_rate":chg["new_rate"]} for chg in sub["changes"]]
                    df_sub,s_sub=build_amortization(b["principal"],b["annual_rate"],b["n_py"],b["amort_years"],accel=b["accel"],start_date=b["start_date"],extra_payments=b.get("past_extra") or None,rate_changes=sub_rc or None)
                    if df_sub.empty: continue
                    sub_id=f"{sc['name']}{sub_k}"
                    rem_sub=round((len(df_sub)-today_p_sc)/b["n_py"],1) if today_p_sc>0 else b["amort_years"]
                    st.markdown(f"**Sub-scenario {sub_id}: {sub['name']}**")
                    s1,s2,s3=st.columns(3)
                    s1.metric("Interest",f"${s_sub.get('total_interest',0):,.0f}",delta=f"${s_sub.get('total_interest',0)-s_base_ref.get('total_interest',0):+,.0f}")
                    s2.metric("Remaining",f"{rem_sub:.1f} yrs")
                    s3.metric("End Balance",f"${s_sub.get('end_balance',0):,.0f}")
                    fig_sub_=stacked_bar_pi(df_sub,today_p_sc,sc_term_end_p,f"Sub {sub_id}: {sub['name']}")
                    st.plotly_chart(fig_sub_,use_container_width=True,key=f"ch_sub_{sc_id}_{rn['id']}_{sub_k}")

            # Save/update scenario
            sav_n=st.text_input("Save as",sc["name"],key=f"sc_save_name_{sc_id}")
            sc_col1,sc_col2=st.columns(2)
            if sc_col1.button("💾 Save scenario",key=f"save_rc_{sc_id}"):
                if not sav_n or not sav_n.strip():
                    sc_col2.error("❌ Scenario name cannot be empty. Please enter a name in the 'Save as' field.")
                elif not sc["renewals"]:
                    sc_col2.error("❌ Add at least one renewal entry before saving.")
                else:
                    sc_params={**b,"start_date":str(b["start_date"]),"rate_changes":main_rc,"sc_name":sc["name"],"sc_desc":sc["desc"]}
                    sc_summary=s_sc
                    existing=db_sc_by_name.get(sav_n.strip())
                    if existing:
                        ok=db_update_scenario(st.session_state.db_conn,existing["id"],sav_n.strip(),sc_params,sc_summary)
                        sc_col2.success(f"✅ Updated in DB: {sav_n}") if ok else sc_col2.error("❌ Update failed — check DB connection.")
                    else:
                        ok=db_save_scenario(st.session_state.db_conn,sav_n.strip(),sc_params,sc_summary)
                        sc_col2.success(f"✅ Saved to DB: {sav_n}") if ok else sc_col2.error("❌ Save failed — check DB connection.")

    for sc_id in sc_del: del rcs[sc_id]
    if sc_del: st.rerun()

    # FIX 4: No separate saved section — DB scenarios shown as collapsable entries
    # alongside newly created ones above.
    db_scenarios_fresh=db_load_scenarios(st.session_state.db_conn)
    if db_scenarios_fresh:
        st.markdown("---")
        st.markdown("##### 📂 Previously Saved Scenarios (click to view / edit / delete)")
        for sc_db in db_scenarios_fresh:
            # Skip if currently being edited (already in rcs)
            already_editing=any(s["name"]==sc_db["name"] for s in rcs.values())
            if already_editing: continue
            s=sc_db["summary"]
            with st.expander(f"💾 {sc_db['name']}  ·  saved {sc_db['created_at'][:16]}",expanded=False):
                cc_=st.columns(4)
                cc_[0].metric("Scenario Interest",f"${s.get('total_interest',0):,.0f}",help="Total interest under this scenario")
                cc_[1].metric("Payoff",f"{s.get('payoff_years',0):.1f} yrs",help="Full amortization payoff period")
                cc_[2].metric("Payment",f"${s.get('payment',0):,.2f}",help="Regular payment amount")
                cc_[3].metric("End Balance",f"${s.get('end_balance',0):,.0f}",help="Balance at end of amortization")
                ea,eb,ec=st.columns(3)
                if ea.button("✏️ Load for Editing",key=f"edit_db_{sc_db['id']}"):
                    rcs_list=sc_db["params"].get("rate_changes",[])
                    nid=str(uuid.uuid4())[:8]
                    rcs[nid]={"name":sc_db["name"],"desc":sc_db["params"].get("sc_desc",""),"renewals":[
                        {"id":str(uuid.uuid4())[:8],"mode":"By Period","date_str":str(period_to_date(rc["period"],b["start_date"],b["n_py"])),
                         "period":rc["period"],"new_rate":rc["new_rate"],"mtype":"Fixed","term_years":3,
                         "actual_penalty":0,"misc_fees":250,"orig_posted":rc["new_rate"]+1.5,"curr_posted":rc["new_rate"]-0.5,"variable_subs":{}}
                        for rc in rcs_list]}
                    st.success(f"Loaded '{sc_db['name']}' into editor above — scroll up."); st.rerun()
                if eb.button("🗑️ Delete",key=f"del_db_{sc_db['id']}"):
                    db_delete_scenario(st.session_state.db_conn,sc_db["id"]); st.success("Deleted."); st.rerun()
                show_raw=ec.checkbox("Raw params",False,key=f"raw_{sc_db['id']}")
                if show_raw: st.json(sc_db.get("params",{}))

    # Education section
    st.divider()
    with st.expander("📚 Canadian Mortgage Education"):
        st.markdown("""**Semi-annual compounding** (Interest Act): Rate 5.39% → Eff. annual `(1+0.0539/2)²=5.463%` · Monthly `1.05463^(1/12)-1=0.4453%`

**CMHC**: <10%=4%, 10-15%=3.1%, 15-20%=2.8%, ≥20%=0%. Price ≤$1.5M required.

**Break penalty**: Variable=3 months interest · Fixed=max(3 months interest, IRD) where IRD=(orig posted−curr posted)×bal×remaining yrs""")

# ══════════════════════════════════════════════════════════════════
# TAB 3 — AMORTIZATION SCHEDULE  (FIX #5, #12)
# ══════════════════════════════════════════════════════════════════
with tabs[2]:
    st.subheader("📅 Full Amortization Schedule")
    require_setup()
    b=st.session_state["base"]
    today_ym=date.today().strftime("%Y-%m")

    # Scenario selector
    db_sc_list=db_load_scenarios(st.session_state.db_conn)
    sc_opts=["Current Setup (base rates)"]+[s["name"] for s in db_sc_list]
    chosen_sc=st.selectbox("Display schedule for:",sc_opts,key="sch_sc_sel",help="Choose a saved scenario to see how the schedule changes")
    if chosen_sc=="Current Setup (base rates)": sc_rcs=b.get("past_renewal_rcs") or []
    else:
        saved_p=next((s["params"] for s in db_sc_list if s["name"]==chosen_sc),{})
        sc_rcs=(b.get("past_renewal_rcs") or [])+(saved_p.get("rate_changes") or [])

    # FIX #5: Build schedule — ends at actual remaining amortization from today
    # Full schedule (all periods from start until payoff)
    df_sch_full,_=build_amortization(b["principal"],b["annual_rate"],b["n_py"],b["amort_years"],accel=b["accel"],start_date=b["start_date"],extra_payments=b.get("past_extra") or None,rate_changes=sc_rcs or None)

    def get_ym(d): return d.strftime("%Y-%m") if hasattr(d,"strftime") else str(d)[:7]
    def tod_d(v): return v if isinstance(v,date) else (v.date() if hasattr(v,"date") else date.fromisoformat(str(v)[:10]))
    today_rows=df_sch_full[df_sch_full["Date"].apply(get_ym)==today_ym]
    today_period=int(today_rows["Period"].iloc[0]) if not today_rows.empty else None

    # FIX #12: Hierarchy view toggle
    view_mode=st.selectbox("View Mode",["Hierarchy (default)","All Periods","Monthly Summary","Yearly Summary"],key="sch_view_mode",help="Hierarchy groups into Past / Current (±4 months) / Future")
    do_hl=st.checkbox("Highlight current month",True,key="sch_hl")

    if view_mode=="Hierarchy (default)" and today_period:
        # Build three segments
        past_df=df_sch_full[df_sch_full["Period"]<max(1,today_period-4)]
        current_df=df_sch_full[(df_sch_full["Period"]>=max(1,today_period-4))&(df_sch_full["Period"]<=today_period)]
        future_df=df_sch_full[df_sch_full["Period"]>today_period]

        def fmt_seg(seg_df, seg_label, is_current=False):
            if seg_df.empty: return
            disp=seg_df[["Period","Date","Payment","Interest","Principal","Balance","Rate (%)","Cum Interest"]].copy()
            disp["Date"]=disp["Date"].apply(lambda d: d.strftime("%b %Y") if hasattr(d,"strftime") else str(d)[:7])
            mc_=[c for c in disp.columns if c not in ["Period","Date","Rate (%)"]]
            # Highlight today in current segment
            def _hl(row):
                if do_hl and is_current and str(row.get("Date",""))[:7]==today_ym:
                    return["background:#FFF3CD;font-weight:bold"]*len(row)
                return[""]*len(row)

            # By Year within segment
            for yr, yr_df in disp.groupby(seg_df.loc[disp.index if not disp.empty else seg_df.index,"CalYear"] if not seg_df.empty else disp["Date"].str[:4]):
                with st.expander(f"📅 {yr} — {len(yr_df)} payment(s)",expanded=is_current):
                    st.dataframe(yr_df.style.apply(_hl,axis=1).format({c:"${:,.2f}" for c in mc_}),use_container_width=True,height=min(60+len(yr_df)*35,350))

        # Attach CalYear back to sub-dfs
        past_df=past_df.copy(); past_df["CalYear"]=past_df["Date"].apply(_year_of) if not past_df.empty else past_df["CalYear"]
        current_df=current_df.copy()
        future_df=future_df.copy(); future_df["CalYear"]=future_df["Date"].apply(_year_of) if not future_df.empty else future_df.get("CalYear",future_df["Period"])

        # Past segment — show as single table grouped by year (no nested expander)
        with st.expander(f"◀ PAST — {len(past_df)} payments (before recent history)",expanded=False):
            if not past_df.empty:
                d2=past_df[["CalYear","Period","Date","Payment","Interest","Principal","Balance","Rate (%)"]].copy()
                d2["Date"]=d2["Date"].apply(lambda d: d.strftime("%b %Y") if hasattr(d,"strftime") else str(d)[:7])
                d2=d2.rename(columns={"CalYear":"Year"})
                mc_=[c for c in d2.columns if c not in ["Year","Period","Date","Rate (%)"]]
                st.dataframe(d2.style.format({c:"${:,.2f}" for c in mc_}),use_container_width=True,height=min(80+len(d2)*28,420))

        # Current segment (expanded by default)
        bal_today=float(df_sch_full[df_sch_full["Period"]==today_period]["Balance"].iloc[0]) if today_period else 0
        with st.expander(f"★ CURRENT — last 4 months + today ({date.today().strftime('%B %Y')}) · Balance: ${bal_today:,.0f}",expanded=True):
            if not current_df.empty:
                d3=current_df[["Period","Date","Payment","Interest","Principal","Balance","Rate (%)","Cum Interest"]].copy()
                d3["Date"]=d3["Date"].apply(lambda d: d.strftime("%b %Y") if hasattr(d,"strftime") else str(d)[:7])
                mc_=[c for c in d3.columns if c not in ["Period","Date","Rate (%)"]]
                def _hl_c(row):
                    if do_hl and str(row.get("Date",""))[:7]==today_ym: return["background:#FFF3CD;font-weight:bold"]*len(row)
                    return[""]*len(row)
                st.dataframe(d3.style.apply(_hl_c,axis=1).format({c:"${:,.2f}" for c in mc_}),use_container_width=True,height=min(60+len(d3)*35,280))

        # Future segment — single table with year column, blue-italic styling
        with st.expander(f"▶ FUTURE — {len(future_df)} remaining payments",expanded=False):
            if not future_df.empty:
                future_df_disp=future_df.copy()
                future_df_disp["CalYear"]=future_df_disp["Date"].apply(_year_of)
                d4=future_df_disp[["CalYear","Period","Date","Payment","Interest","Principal","Balance","Rate (%)"]].copy()
                d4["Date"]=d4["Date"].apply(lambda d: d.strftime("%b %Y") if hasattr(d,"strftime") else str(d)[:7])
                d4=d4.rename(columns={"CalYear":"Year"})
                mc_=[c for c in d4.columns if c not in ["Year","Period","Date","Rate (%)"]]
                def _hl_f(row): return["background:#e8f4fd;font-style:italic;color:#555"]*len(row)
                st.dataframe(d4.style.apply(_hl_f,axis=1).format({c:"${:,.2f}" for c in mc_}),use_container_width=True,height=min(80+len(d4)*28,460))

    else:
        # Regular table modes
        if view_mode=="Yearly Summary":
            disp=df_sch_full.groupby("CalYear").agg(Payments=("Payment","count"),Total_Paid=("Total Paid","sum"),Interest=("Interest","sum"),Principal=("Principal","sum"),Ending_Balance=("Balance","last")).reset_index()
            disp.columns=["Year","Payments","Total Paid","Interest","Principal","Ending Balance"]
            cur_y=date.today().year
            def _hl(row):
                yr=int(row.get("Year",0))
                if do_hl and yr==cur_y: return["background:#FFF3CD;font-weight:bold"]*len(row)
                if yr>cur_y: return["background:#e8f4fd;font-style:italic;color:#666"]*len(row)
                return[""]*len(row)
        elif view_mode=="Monthly Summary" and b["n_py"]>12:
            df_sch_full["YM"]=df_sch_full["Date"].apply(get_ym)
            disp=df_sch_full.groupby("YM").agg(Total_Paid=("Total Paid","sum"),Interest=("Interest","sum"),Principal=("Principal","sum"),Ending_Balance=("Balance","last")).reset_index()
            def _hl(row):
                ym=str(row.get("YM",""))
                if do_hl and ym==today_ym: return["background:#FFF3CD;font-weight:bold"]*len(row)
                if ym>today_ym: return["background:#e8f4fd;font-style:italic;color:#555"]*len(row)
                return[""]*len(row)
        else:
            # Start near today (FIX #5)
            show_all=st.checkbox("Show full schedule from start",False,key="sch_show_all")
            disp=df_sch_full[["Period","Date","Payment","Interest","Principal","Balance","Rate (%)","Cum Interest"]].copy()
            disp["Date"]=disp["Date"].apply(lambda d: d.strftime("%Y-%m-%d") if hasattr(d,"strftime") else str(d)[:10])
            if not show_all and today_period:
                disp=disp.iloc[max(0,today_period-4):].reset_index(drop=True)
            def _hl(row):
                ds=str(row.get("Date",""))[:7]
                if do_hl and ds==today_ym: return["background:#FFF3CD;font-weight:bold"]*len(row)
                if ds>today_ym: return["background:#e8f4fd;font-style:italic;color:#555"]*len(row)
                return[""]*len(row)

        mc_=[c for c in disp.columns if c not in ["Period","Year","Payments","YM","Date","Rate (%)"]]
        st.dataframe(disp.style.apply(_hl,axis=1).format({c:"${:,.2f}" for c in mc_}),use_container_width=True,height=500)

    if today_period:
        bal_t=float(df_sch_full[df_sch_full["Period"]==today_period]["Balance"].iloc[0])
        rem_p=len(df_sch_full)-today_period
        rem_y=round(rem_p/b["n_py"],1)
        st.markdown(f'<div class="ok">🟡 Current: Period <b>{today_period}</b> ({date.today().strftime("%B %Y")}) · Balance: <b>${bal_t:,.0f}</b> · Remaining: <b>{rem_y:.1f} yrs</b> · <i style="color:#555">Blue-italic = future projections</i></div>',unsafe_allow_html=True)

    fig_bal=go.Figure()
    fig_bal.add_scatter(x=df_sch_full["Date"],y=df_sch_full["Balance"],fill="tozeroy",name="Balance",line=dict(color="#1a3c5e"))
    fig_bal.add_scatter(x=df_sch_full["Date"],y=df_sch_full["Cum Interest"],name="Cum Interest",line=dict(color="#e74c3c",dash="dash"))
    if today_period:
        td_d=df_sch_full[df_sch_full["Period"]==today_period]["Date"].iloc[0]
        fig_bal.add_vline(x=_vline_x(td_d),line_dash="dash",line_color="#27ae60",annotation_text="Today",annotation_position="top right")
    fig_bal.update_layout(title=f"Balance & Cumulative Interest — {chosen_sc}",xaxis_title="Date",yaxis_title="($)",height=360,margin=dict(t=60,b=40))
    st.plotly_chart(fig_bal,use_container_width=True,key="ch_sch_bal")
    st.download_button("⬇️ Download CSV",df_sch_full.to_csv(index=False).encode(),"schedule.csv","text/csv")

# ══════════════════════════════════════════════════════════════════
# TAB 4 — PREPAYMENT ANALYSIS
# ══════════════════════════════════════════════════════════════════
with tabs[3]:
    st.subheader("💰 Prepayment Analysis")
    require_setup()
    b=st.session_state["base"]; fn=b["n_py"]

    db_sc_pp=db_load_scenarios(st.session_state.db_conn)
    chosen_rc=st.selectbox("Rate scenario base",["Base Rate — no rate changes"]+[s["name"] for s in db_sc_pp],key="pp_rc_sel",help="Select a saved rate scenario to combine with prepayments")
    if chosen_rc=="Base Rate — no rate changes": active_rc=[]
    else:
        saved_p=next((s["params"] for s in db_sc_pp if s["name"]==chosen_rc),{})
        active_rc=saved_p.get("rate_changes",[])
    all_rcs_pp=(b.get("past_renewal_rcs") or [])+active_rc

    col_pp1,col_pp2=st.columns(2)
    with col_pp1:
        st.markdown("##### 📅 Annual Lump-Sum Prepayments")
        annual_lump=st.number_input("Annual lump-sum ($)",0,500_000,10_000,500,key="pp_al",help="Additional principal paid each year")
        lump_month=st.selectbox("Month each year",["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"],key="pp_lm")
        lump_start=st.number_input("Starting year",1,30,1,key="pp_ls")
        lump_nyrs=st.number_input("For how many years?",1,30,5,key="pp_ln")
        lmm={"Jan":1,"Feb":2,"Mar":3,"Apr":4,"May":5,"Jun":6,"Jul":7,"Aug":8,"Sep":9,"Oct":10,"Nov":11,"Dec":12}
        future_extra=[]
        if annual_lump>0:
            for yr in range(int(lump_start),int(lump_start+lump_nyrs)):
                p=max(1,int((yr-1)*fn+lmm[lump_month]*fn/12))
                future_extra.append({"period":p,"amount":float(annual_lump)})
        pp_lim=st.slider("Lender prepayment limit (%)",10,30,20,key="pp_lim")
        if annual_lump>b["principal"]*pp_lim/100:
            st.warning(f"⚠️ Exceeds {pp_lim}% limit (${b['principal']*pp_lim/100:,.0f}).")

    with col_pp2:
        st.markdown("##### 💳 Increased Regular Payments")
        inc_t=st.radio("Increase type",["Fixed $","% increase","None"],index=2,horizontal=True,key="pp_it")
        inc_v=0.0
        if inc_t=="Fixed $": inc_v=float(st.number_input("Extra/payment ($)",0,10_000,200,50,key="pp_if"))
        elif inc_t=="% increase":
            inc_pct=st.slider("% increase",1,100,10,key="pp_ip")
            inc_v=calc_pmt(b["principal"],b["annual_rate"],fn,b["amort_years"],b["accel"])*inc_pct/100
        if inc_v>0:
            for p in range(1,int(b["amort_years"]*fn)+1): future_extra.append({"period":p,"amount":inc_v})
        st.markdown("##### 🔁 One-Time Lump Sum")
        ot_mode=st.radio("Mode",["By Date","By Period"],horizontal=True,key="pp_om")
        if ot_mode=="By Date":
            ot_d=st.date_input("Date",b["start_date"]+relativedelta(years=1),min_value=b["start_date"],key="pp_od")
            ot_p=date_to_period(ot_d,b["start_date"],fn); st.caption(f"≈ Period {ot_p}")
        else: ot_p=int(st.number_input("Period #",1,int(b["amort_years"]*fn),fn,key="pp_op"))
        ot_a=st.number_input("Amount ($)",0,2_000_000,0,1_000,key="pp_oa")
        if ot_a>0: future_extra.append({"period":int(ot_p),"amount":float(ot_a)})

    past_extra=b.get("past_extra",[]); all_extra=past_extra+future_extra
    df_rsc,s_rsc=build_amortization(b["principal"],b["annual_rate"],b["n_py"],b["amort_years"],accel=b["accel"],start_date=b["start_date"],extra_payments=past_extra or None,rate_changes=all_rcs_pp or None)
    df_pp,s_pp=build_amortization(b["principal"],b["annual_rate"],b["n_py"],b["amort_years"],accel=b["accel"],start_date=b["start_date"],extra_payments=all_extra or None,rate_changes=all_rcs_pp or None)

    tm_pp=b["today_m"]; today_p_pp=tm_pp.get("period_today",0)
    base_rem_pp=tm_pp.get("remaining_years",b["amort_years"])
    rem_rsc=round((len(df_rsc)-today_p_pp)/b["n_py"],1) if today_p_pp>0 and not df_rsc.empty else b["amort_years"]
    rem_pp_v=round((len(df_pp)-today_p_pp)/b["n_py"],1) if today_p_pp>0 and not df_pp.empty else b["amort_years"]
    int_saved=s_rsc.get("total_interest",0)-s_pp.get("total_interest",0)
    new_total=sum(e["amount"] for e in future_extra)

    st.divider()
    m1,m2,m3,m4,m5,m6=st.columns(6)
    m1.metric("Interest (rate sc.)",f"${s_rsc.get('total_interest',0):,.0f}",help="Total interest under selected rate scenario")
    m2.metric("Interest (+ prepayments)",f"${s_pp.get('total_interest',0):,.0f}",delta=f"${-int_saved:+,.0f}",help="Total interest with prepayments")
    m3.metric("Remaining (rate sc.)",f"{rem_rsc:.1f} yrs",help="Remaining amortization from today under rate scenario")
    m4.metric("Remaining (+ prepayments)",f"{rem_pp_v:.1f} yrs",delta=f"{rem_pp_v-rem_rsc:+.1f} yrs",help="Remaining with prepayments")
    m5.metric("Total New Prepaid",f"${new_total:,.0f}",help="Total additional prepayments planned")
    m6.metric("Interest ROI",f"{int_saved/new_total*100:.1f}%" if new_total>0 and int_saved>0 else "—",help="Interest saved per dollar prepaid × 100%")
    if int_saved>0:
        st.markdown(f'<div class="ok">💚 Prepayments save <b>${int_saved:,.0f}</b> · Shorten by <b>{rem_rsc-rem_pp_v:.1f} yrs</b></div>',unsafe_allow_html=True)

    sc_end_p=int(b["term_years"]*b["n_py"])+int(3*b["n_py"])
    if not df_pp.empty:
        fig_pp_bar=stacked_bar_pi(df_pp,today_p_pp,sc_end_p,f"Prepayment Impact — P & I ({chosen_rc})")
        st.plotly_chart(fig_pp_bar,use_container_width=True,key="ch_pp_bar")

    sc_np=st.text_input("Save as","Prepayment Scenario",key="pp_scname")
    if st.button("💾 Save",key="btn_save_pp"):
        sc_p={"rate_changes":active_rc,"extra_payments":len(all_extra),"sc_type":"prepayment"}
        ok=db_save_scenario(st.session_state.db_conn,sc_np,sc_p,s_pp)
        st.success("Saved to DB" if ok else "Save failed")

# ══════════════════════════════════════════════════════════════════
# TAB 5 — BREAK PENALTY
# ══════════════════════════════════════════════════════════════════
with tabs[4]:
    st.subheader("⚠️ Mortgage Break Penalty Calculator")
    require_setup()
    b=st.session_state["base"]
    c1,c2=st.columns(2)
    with c1:
        bp_bal=st.number_input("Outstanding Balance ($)",100,5_000_000,int(b.get("principal",500_000)*0.85),1_000,key="bp_bal",help="Current outstanding mortgage balance")
        bp_rate=st.number_input("Contract Rate (%)",0.5,20.0,float(b.get("annual_rate",5.39)),0.01,format="%.2f",key="bp_rate",help="Your current mortgage interest rate")
        bp_mtype=st.selectbox("Mortgage Type",["Fixed","Variable"],index=0 if b.get("mortgage_type","Fixed")=="Fixed" else 1,key="bp_mtype_tab5",help="Fixed or Variable rate")
        bp_mleft=st.slider("Months Remaining in Term",1,120,36,key="bp_mleft",help="Months left in current term")
        bp_misc=st.number_input("Miscellaneous Fees ($)",0,50_000,500,50,key="bp_misc",help="Admin, appraisal, legal fees")
    with c2:
        if bp_mtype=="Fixed":
            st.markdown("##### IRD Inputs")
            bp_orig=st.number_input("Posted Rate at Origination (%)",0.5,20.0,float(b.get("annual_rate",5.39))+1.5,0.01,format="%.2f",key="bp_orig",help="Bank's posted rate when you originally signed")
            bp_curr=st.number_input("Current Posted Rate for Remaining Term (%)",0.5,20.0,max(float(b.get("annual_rate",5.39))-0.5,0.5),0.01,format="%.2f",key="bp_curr",help="Current posted rate for remaining term length")
        else: bp_orig=bp_curr=float(b.get("annual_rate",5.39))

    pen=calc_break_penalty(bp_bal,bp_rate,bp_mtype,bp_orig,bp_curr,bp_mleft)
    pen_opts=[f"3-Month Interest (${pen['3_months_interest']:,.0f})"]
    if pen["ird"] is not None: pen_opts.append(f"IRD (${pen['ird']:,.0f})")
    pen_opts.append("Custom value")
    bpc1,bpc2=st.columns([3,2])
    bp_choice=bpc1.radio("Apply which penalty?",pen_opts,key="bp_pen_radio")
    if "Custom" in bp_choice:
        bp_custom_str=bpc2.text_input("",value=str(int(pen["calc_penalty"])),key="bp_custom_inp")
        bpc2.caption("Custom value ($)")
        try: actual_pen=float(bp_custom_str.replace(",","").replace("$",""))
        except: actual_pen=pen["calc_penalty"]
    elif "IRD" in bp_choice: actual_pen=pen["ird"] or 0.0
    else: actual_pen=pen["3_months_interest"]

    total_exit=actual_pen+bp_misc
    cc1,cc2,cc3,cc4=st.columns(4)
    cc1.metric("3 Months Interest",f"${pen['3_months_interest']:,.2f}",help="3 months interest on outstanding balance")
    if pen["ird"] is not None: cc2.metric("IRD",f"${pen['ird']:,.2f}",help="Interest Rate Differential")
    cc3.metric("Penalty Applied",f"${actual_pen:,.2f}",help="The penalty amount you chose")
    cc4.metric("Total Exit Cost",f"${total_exit:,.2f}",help="Penalty + misc fees")
    st.divider()
    new_r=st.slider("New rate if you break (%)",0.5,15.0,max(float(b.get("annual_rate",5.39))-1.0,0.5),0.05,key="bp_newr")
    ar=max(bp_mleft/12,1)
    _,s_stay=build_amortization(bp_bal,bp_rate,12,ar,term_periods=bp_mleft)
    _,s_brk =build_amortization(bp_bal,new_r,12,ar,term_periods=bp_mleft)
    int_stay=s_stay.get("total_interest",0); int_brk=s_brk.get("total_interest",0)+total_exit
    bc1,bc2,bc3=st.columns(3)
    bc1.metric("Interest (Stay)",f"${int_stay:,.0f}",help="Interest over remaining term if you stay")
    bc2.metric("Interest+Fees (Break)",f"${int_brk:,.0f}",help="Interest at new rate + exit fees")
    bc3.metric("Net Savings",f"${int_stay-int_brk:,.0f}",delta="✅ Worth breaking" if int_stay>int_brk else "❌ Not worth it",help="Positive = breaking saves money")
    sweep=np.arange(0.5,float(b.get("annual_rate",5.39))+0.11,0.25)
    svlist=[]
    for tr in sweep:
        _,st_=build_amortization(bp_bal,tr,12,ar,term_periods=bp_mleft)
        svlist.append(int_stay-(st_.get("total_interest",0)+total_exit))
    fig_be=go.Figure()
    fig_be.add_scatter(x=list(sweep),y=svlist,mode="lines+markers",line=dict(color="#1a3c5e"),name="Net Savings")
    fig_be.add_hline(y=0,line_dash="dash",line_color="red",annotation_text="Break-even",annotation_position="top right")
    fig_be.add_vline(x=float(b.get("annual_rate",5.39)),line_dash="dot",line_color="orange",annotation_text="Current rate",annotation_position="top left")
    fig_be.update_layout(title="Net Savings vs New Rate",xaxis_title="New Rate (%)",yaxis_title="Net Savings ($)",height=320,margin=dict(t=50,b=40))
    st.plotly_chart(fig_be,use_container_width=True,key="ch_bpbe")
    op=calc_pmt(bp_bal,bp_rate,12,ar); np_=calc_pmt(bp_bal,new_r,12,ar)
    if abs(op-np_)>1:
        st.markdown(f'<div class="inf">Monthly: <b>${op:,.2f}</b> → <b>${np_:,.2f}</b> (<b>${np_-op:+,.2f}/mo</b>) · Recoup in: <b>{total_exit/abs(op-np_):.0f} months</b></div>',unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════
# TAB 6 — SCENARIO COMPARISON  (FIX #9: includes current mortgage)
# ══════════════════════════════════════════════════════════════════
with tabs[5]:
    st.subheader("🔄 Side-by-Side Scenario Comparison")
    require_setup()
    b=st.session_state["base"]

    n_sc=st.radio("Number of scenarios to compare",[2,3,4],horizontal=True,key="cmp_n")
    db_sc_cmp=db_load_scenarios(st.session_state.db_conn)
    # FIX #9: Scenario options include Current Mortgage + all saved scenarios
    sc_option_names=["Current Mortgage (base)"]+[s["name"] for s in db_sc_cmp]

    sc_defs=[]; cols=st.columns(int(n_sc))
    for i,col in enumerate(cols):
        with col:
            st.markdown(f"**Scenario {i+1}**")
            # First scenario defaults to Current Mortgage
            default_opt=0 if i==0 else min(i,len(sc_option_names)-1)
            sc_pick=st.selectbox(f"Scenario {i+1} source",sc_option_names,index=default_opt,key=f"cmp_src_{i}",help="Choose current mortgage or a saved scenario")

            if sc_pick=="Current Mortgage (base)":
                rate=float(b["annual_rate"]); amt=b["amort_years"]; frq=b["payment_freq"]
                lbl="Current Mortgage"
                rc_list=b.get("past_renewal_rcs") or []
            else:
                sc_match=next((s for s in db_sc_cmp if s["name"]==sc_pick),{})
                sc_par=sc_match.get("params",{})
                rate=float(sc_par.get("sc_rate",b["annual_rate"]) or sc_par.get("rate_changes",[{}])[-1].get("new_rate",b["annual_rate"]) if sc_par.get("rate_changes") else b["annual_rate"])
                amt=sc_par.get("amort_years",b["amort_years"])
                frq=sc_par.get("payment_freq",b["payment_freq"])
                rc_list=(b.get("past_renewal_rcs") or [])+(sc_par.get("rate_changes") or [])
                lbl=sc_pick

            lbl=st.text_input("Label",lbl,key=f"cmp_lbl_{i}")
            rate=st.number_input("Rate (%)",0.5,20.0,float(rate),0.01,key=f"cmp_rate_{i}",format="%.2f")
            amt=st.slider("Amort (yrs)",5,30,int(amt),key=f"cmp_amt_{i}")
            frq=st.selectbox("Frequency",list(FREQ.keys()),index=list(FREQ.keys()).index(b["payment_freq"]),key=f"cmp_frq_{i}")
            lump=st.number_input("Annual lump ($)",0,200_000,0,1_000,key=f"cmp_lump_{i}")
            fc_=FREQ[frq]; ny=fc_["n"]; ac=fc_["accel"]
            ex=list(b.get("past_extra",[]))
            if lump>0:
                for yr in range(1,amt+1): ex.append({"period":max(1,int((yr-1)*ny+ny//2)),"amount":float(lump)})
            df_c,s_c=build_amortization(b["principal"],rate,ny,amt,accel=ac,start_date=b["start_date"],extra_payments=ex or None,rate_changes=rc_list or None)
            pmt_c=calc_pmt(b["principal"],rate,ny,amt,ac)
            tp_c=b["today_m"].get("period_today",0)
            rem_c=round((len(df_c)-tp_c)/ny,1) if tp_c>0 and not df_c.empty else amt
            today_bal_c=b["today_m"].get("balance_today",b["principal"])
            pmt_today_c=calc_pmt(today_bal_c,rate,ny,rem_c,ac) if rem_c>0 else pmt_c
            sc_defs.append({"label":lbl,"rate":rate,"amort":amt,"freq":frq,"lump":lump,"df":df_c,"summary":s_c,"payment":pmt_c,"pmt_today":pmt_today_c,"rem":rem_c,"n_py":ny})

    st.divider()
    comp_rows=[]
    for sc in sc_defs:
        s=sc["summary"]
        comp_rows.append({"Scenario":sc["label"],"Rate":f"{sc['rate']:.2f}%","Amort":f"{sc['amort']} yrs","Frequency":sc["freq"],"Annual Lump":f"${sc['lump']:,.0f}","Orig Payment":f"${sc['payment']:,.2f}","Current Payment":f"${sc['pmt_today']:,.2f}","Remaining":f"{sc['rem']:.1f} yrs","Total Interest":f"${s.get('total_interest',0):,.0f}","Total Paid":f"${s.get('total_paid',0):,.0f}"})
    st.dataframe(pd.DataFrame(comp_rows),use_container_width=True)

    pal=["#1a3c5e","#e74c3c","#27ae60","#f39c12"]
    fig_c=go.Figure()
    for i,sc in enumerate(sc_defs):
        if not sc["df"].empty:
            fig_c.add_scatter(x=sc["df"]["Date"],y=sc["df"]["Balance"],name=sc["label"],line=dict(color=pal[i]))
    # FIX #11: annotation for today above chart area
    if b["today_m"].get("period_today"):
        td_d2=df_sch_full["Date"].iloc[b["today_m"]["period_today"]-1] if not df_sch_full.empty and b["today_m"]["period_today"]<=len(df_sch_full) else None
        if td_d2 is not None:
            fig_c.add_vline(x=_vline_x(td_d2),line_dash="dash",line_color="#27ae60",annotation_text="Today",annotation_position="top right")
    fig_c.update_layout(title="Balance Comparison",xaxis_title="Date",yaxis_title="($)",height=340,margin=dict(t=60,b=40))
    st.plotly_chart(fig_c,use_container_width=True,key="ch_cmpbal")

    fig_cmp_bar=go.Figure()
    for i,sc in enumerate(sc_defs):
        if sc["df"].empty: continue
        g=sc["df"].groupby("CalYear").agg(Principal=("Principal","sum"),Interest=("Interest","sum")).reset_index()
        fig_cmp_bar.add_bar(x=g["CalYear"].astype(str),y=g["Principal"],name=f"{sc['label']} P",marker_color=pal[i],opacity=0.9,legendgroup=sc["label"])
        fig_cmp_bar.add_bar(x=g["CalYear"].astype(str),y=g["Interest"],name=f"{sc['label']} I",marker_color=pal[i],opacity=0.5,legendgroup=sc["label"])
    fig_cmp_bar.update_layout(barmode="stack",title="Annual P & I by Scenario",xaxis_title="Year",yaxis_title="($)",height=360,margin=dict(t=70,b=40),legend=dict(orientation="h",yanchor="bottom",y=1.02))
    st.plotly_chart(fig_cmp_bar,use_container_width=True,key="ch_cmpbar")

    best=min(range(len(sc_defs)),key=lambda i:sc_defs[i]["summary"].get("total_interest",1e12))
    worst_i=max(sc["summary"].get("total_interest",0) for sc in sc_defs)
    st.markdown(f'<div class="ok">🏆 <b>{sc_defs[best]["label"]}</b> saves ${worst_i-sc_defs[best]["summary"].get("total_interest",0):,.0f} · Remaining: <b>{sc_defs[best]["rem"]:.1f} yrs</b></div>',unsafe_allow_html=True)

    # Prompt download
    st.divider()
    PROMPT="""Build Canadian Mortgage Analyzer Streamlit app (app.py). MANDATORY MS SQL Server via pyodbc.
DEFAULTS: DB=localhost\\SQLEXPRESS/MortgageDB/Windows Auth. Mortgage: $1,030,000 | 20% down | 5.39% | 30yr amort | 3yr term | 2023-08-15 | Monthly | Fixed.
MATH: periodic_rate uses two-step: eff=((1+r/200)**2); return eff**(1/n)-1. Build CalYear column in df.
DB: mortgage_setup (single row), mortgage_scenarios(id,name,created_at,params,summary). DB-ONLY scenarios (no local duplicates).
TAB ORDER: Setup | Rate Change Scenarios | Amortization Schedule | Prepayment | Break Penalty | Comparison.
SETUP: Split into sections A(Purchase/Down) B(Mortgage Terms) C(Past Renewals collapsable) D(Past Prepayments collapsable).
KEY METRICS: Initial Principal, Balance@TermEnd, Balance Today, Principal Paid, Interest Paid, Current Remaining Amortization (yrs + end date), Total Interest, Original Amortization Period, Current Monthly Payment. NO "Original Payoff" metric.
STACKED BAR: stacked_bar_pi(df,today_p,term_end_p,title): x=CalYear strings, 3 colour segments (past=grey, current=blue/red, post=faded). Annotations at y=1.12/1.08 paper coords to avoid legend overlap. No add_vline on categorical axis.
AMORTIZATION: Hierarchy toggle default: Past/Current(last4+today)/Future segments, each with Year sub-expanders. Full schedule ends at actual payoff period (len(df)).
SCENARIO METRICS: base_remaining = today_m["remaining_years"] not amort_years.
COMPARISON: first scenario defaults to "Current Mortgage (base)"; all saved DB scenarios selectable.
WORD WIREFRAME: generate_wireframe_docx() using python-docx creating layout reference without any financial numbers. Download button on setup page.
PENALTY: radio 3-Month/IRD/Custom + inline text_input for custom.
ALL METRICS: help= tooltips. ALL CHARTS: CalYear string or date x-axis (never period numbers). FIX #11: annotations pushed to y>=1.08 paper coords.
RUN: streamlit run app.py"""
    st.download_button("📥 Download Fresh-Chat Prompt (.txt)",data=PROMPT.encode("utf-8"),file_name="mortgage_analyzer_prompt.txt",mime="text/plain",key="btn_dl_prompt")
